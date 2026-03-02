#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
DOCX嵌入式Excel表格处理工具 - 多功能单文件CLI工具

程序名称: docx-allinone.py (v13 - 完整版)
作者: BlueSkyXN
描述: 智能处理DOCX文档中的嵌入式Excel表格，支持多种转换模式和文档优化功能

========================================
核心功能
========================================

1. Excel表格处理（多模式可组合）:
   --word-table:      将Excel表格转换为可编辑的Word原生表格（默认模式）
   --extract-excel:   提取嵌入的Excel文件为独立.xlsx文件，并在原位置标注"表 X"
   --image:           将Excel表格渲染成高质量图片（matplotlib引擎，300DPI）
   --keep-attachment: 保留Excel附件入口，用户仍可双击查看完整表格

2. 文档优化功能:
   --remove-watermark: 智能移除各类水印（文本、图片、背景、VML水印）
   --a3:              设置A3横向页面布局（420mm×297mm），优化大表格显示

========================================
技术特性
========================================

• 精准Excel识别: 基于VML+OLE+ProgID多重特征识别，支持多种Excel嵌入格式
• 原表格保护:    完全不影响用户手动创建的Word表格，仅处理嵌入Excel对象
• 智能样式保留:  保留Excel单元格的字体、颜色、对齐、合并、格式等完整样式
• 自适应尺寸:    列宽自动调整(1.5-6.0cm)，行高基于内容(≥0.6cm)
• 中文字体支持:  自动检测系统字体（macOS: Heiti TC, Windows: SimHei, Linux: WenQuanYi）
• 两阶段处理:    先插入新内容，后智能清理，确保文档结构完整性
• 逆序处理:      从后往前处理嵌入对象，避免删除时索引错位

========================================
水印移除能力
========================================

DocumentCleaner类支持移除：
- VML文本水印（shape + textpath结构）
- 图片水印（居中定位 + behindDoc属性）
- 背景水印（background元素）
- Word Art对象
- 页眉页脚中的所有水印类型

========================================
依赖库
========================================

pip install python-docx openpyxl pillow matplotlib

========================================
使用示例
========================================

# 默认模式 (转为Word表格)
python docx-allinone.py document.docx

# 将表格转为高清图片
python docx-allinone.py document.docx --image

# 提取Excel文件 + 插入图片
python docx-allinone.py document.docx --extract-excel --image

# 全功能模式（表格+提取+图片）
python docx-allinone.py document.docx --word-table --extract-excel --image

# 保留附件 + 转换表格
python docx-allinone.py document.docx --word-table --keep-attachment

# 移除水印 + A3横向布局
python docx-allinone.py document.docx --remove-watermark --a3

# 完整优化（A3横向 + 图片 + 无水印）
python docx-allinone.py document.docx --image --a3 --remove-watermark

========================================
输出文件命名规则
========================================

基础名称: [原文件名]-AIO.docx
附加后缀:
  - WithAttachments: 保留了Excel附件
  - NoWM:           移除了水印
  - A3:             A3横向布局

示例: document-AIO-A3-NoWM.docx

========================================
技术说明
========================================

• Excel对象识别: 使用XML解析检测VML shape、OLE对象、ProgID等特征
• 关系ID匹配:    通过document.xml.rels文件匹配Excel嵌入文件
• 段落保护:      检测段落父节点，避免误处理表格内的嵌入对象
• 格式化处理:    支持Excel的number_format（百分比、小数、整数等）
• 合并单元格:    使用openpyxl的merged_cells.ranges信息准确还原
"""

import os
import sys
import argparse
import zipfile
import io
import glob
import tempfile
import traceback
import contextlib
import concurrent.futures
import csv
from pathlib import Path
import openpyxl
import xml.etree.ElementTree as ET
import matplotlib
matplotlib.use('Agg')  # 设置非交互式后端，避免GUI依赖
import matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement, parse_xml

# 可选依赖：reportlab + PIL（仅 --table-extract / --split-images 需要）
# 延迟导入，缺失时核心模式（--word-table/--image/--remove-watermark/--a3）仍可正常使用
_HAS_REPORTLAB = False
_HAS_PIL = False
try:
    from reportlab.lib.pagesizes import A4, A3
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        BaseDocTemplate, PageTemplate, Frame, LongTable, TableStyle,
        Paragraph, Spacer, PageBreak, NextPageTemplate, Flowable
    )
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT
    _HAS_REPORTLAB = True
except ImportError:
    pass

try:
    from PIL import Image
    _HAS_PIL = True
except ImportError:
    pass


class DocumentSkipError(Exception):
    """文档处理被跳过（非异常，表示无需处理）"""
    pass


# --- 常量定义 ---

CAPTION_KEYWORDS = [
    "点击图片可查看完整电子表格",
    "Click to view complete spreadsheet",
    "查看完整电子表格"
]

# A3纸张尺寸常量 (单位：英寸)
A3_WIDTH_LANDSCAPE = 16.54  # 420mm
A3_HEIGHT_LANDSCAPE = 11.69  # 297mm

# 输出文件标签（用于过滤已处理的文件）
OUTPUT_FILE_TAGS = [
    "-WithAttachments",
    "-NoWM",
    "-A3",
    "-AIO",
    "_已标记图片",
    "_已标记表格"
]

# 表格提取输出标记
TABLE_MARK_SUFFIX = "_已标记表格"

# 图片分离PDF页面尺寸（点）— reportlab 可用时从 A4/A3 获取，否则使用标准值
if _HAS_REPORTLAB:
    IMG_A4_WIDTH, IMG_A4_HEIGHT = A4
    IMG_A3_WIDTH, IMG_A3_HEIGHT = A3
else:
    IMG_A4_WIDTH, IMG_A4_HEIGHT = 595.276, 841.89
    IMG_A3_WIDTH, IMG_A3_HEIGHT = 841.89, 1190.55

# 注册 VML 命名空间，兼容旧版图片结构
if 'v' not in nsmap:
    nsmap['v'] = 'urn:schemas-microsoft-com:vml'
if 'o' not in nsmap:
    nsmap['o'] = 'urn:schemas-microsoft-com:office:office'


# --- 背景和水印处理功能 ---

class DocumentCleaner:
    """文档水印清理器

    功能说明:
        智能检测并移除Word文档中的各类水印元素，包括页眉页脚、正文和背景中的水印。

    支持的水印类型:
        1. VML文本水印: 使用VML shape + textpath结构的文本水印（最常见）
        2. 图片水印: 使用wp:anchor定位的图片水印（behindDoc + 居中定位）
        3. 背景水印: 文档级背景水印（w:background元素）
        4. Word Art对象: 艺术字水印
        5. VML图片段落: 包含VML shape的w:pict段落

    识别策略:
        - 关键词匹配: 水印文本或ID包含预定义关键词（机密、水印、draft等）
        - 样式特征: 绝对定位 + 旋转 + 居中 + z-index负值
        - 结构特征: behindDoc属性 + 居中对齐
        - ID模式: PowerPlusWaterMark、WordPictureWatermark等

    处理流程:
        1. 分析所有section的页眉页脚（包括首页、偶数页）
        2. 扫描正文中的水印元素
        3. 移除文档级背景水印
        4. 统计并返回处理结果

    Args:
        document: python-docx Document对象

    Attributes:
        document: Word文档对象
        namespaces: XML命名空间映射字典
        watermark_keywords: 水印识别关键词列表（中英文）
    """

    def __init__(self, document):
        self.document = document
        self.namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            've': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
            'o': 'http://schemas.microsoft.com/office/office',
            'r_id': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'v': 'urn:schemas-microsoft-com:vml',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
        }
        
        # 水印识别关键词
        self.watermark_keywords = [
            # 中文水印
            '机密', '水印', '草稿', '样本', '副本', '内部资料', '保密', '绝密',
            # 英文水印
            'confidential', 'watermark', 'draft', 'sample', 'copy', 'internal',
            'secret', 'classified', 'private', 'restricted',
            # 常见水印ID模式
            'powerpluswatermark', 'watermarkobject', 'wordpicturewatermark',
            # Word标准水印ID
            'picturewater', 'waterpicture', '_watermark_', 'wmobj'
        ]
    
    def remove_watermarks(self):
        """移除文档水印 - 增强版"""
        print("    🧹 移除文档水印...")
        removed_count = 0
        
        try:
            # 首先分析水印
            watermark_analysis = self._analyze_all_watermarks()
            print(f"        🔍 检测到水印: 文本={len(watermark_analysis['text'])} 图片={len(watermark_analysis['image'])} 背景={len(watermark_analysis['background'])}")
            
            # 处理每个section的页眉页脚
            for section_idx, section in enumerate(self.document.sections):
                try:
                    # 获取所有可能的页眉页脚
                    headers_footers = [
                        ('header', section.header),
                        ('footer', section.footer),
                        ('first_page_header', getattr(section, 'first_page_header', None)),
                        ('first_page_footer', getattr(section, 'first_page_footer', None)),
                        ('even_page_header', getattr(section, 'even_page_header', None)),
                        ('even_page_footer', getattr(section, 'even_page_footer', None)),
                    ]
                    
                    for hf_name, hf_element in headers_footers:
                        if hf_element:
                            count = self._remove_watermarks_from_header_footer(hf_element, f"Section{section_idx+1}.{hf_name}")
                            removed_count += count
                        
                except Exception as e:
                    print(f"        ⚠️  清理section {section_idx+1} 水印失败: {e}")
            
            # 移除文档正文中的水印元素
            body_count = self._remove_watermarks_from_body()
            removed_count += body_count
            
            # 移除文档级背景水印
            bg_count = self._remove_background_watermarks()
            removed_count += bg_count
            
            print(f"        ✅ 移除水印元素: {removed_count} 个")
            return removed_count > 0
            
        except Exception as e:
            print(f"        ❌ 水印移除失败: {e}")
            return False
    
    def _analyze_all_watermarks(self):
        """分析文档中的所有水印"""
        analysis = {'text': [], 'image': [], 'background': []}
        
        # 分析页眉页脚
        for section in self.document.sections:
            for hf in [section.header, section.footer]:
                if hf:
                    element_analysis = self._analyze_element_watermarks(hf._element)
                    for key in analysis:
                        analysis[key].extend(element_analysis[key])
        
        # 分析正文
        if self.document.element.body:
            body_analysis = self._analyze_element_watermarks(self.document.element.body)
            for key in analysis:
                analysis[key].extend(body_analysis[key])
        
        return analysis
    
    def _analyze_element_watermarks(self, element):
        """分析元素中的水印"""
        watermarks = {'text': [], 'image': [], 'background': []}
        
        try:
            # 检测VML文本水印 - 使用findall替代xpath
            for shape in element.findall('.//{urn:schemas-microsoft-com:vml}shape'):
                # 检查文本路径
                textpaths = shape.findall('.//{urn:schemas-microsoft-com:vml}textpath')
                if textpaths:
                    text_content = textpaths[0].get('string', '').lower()
                    style = shape.get('style', '').lower()
                    shape_id = shape.get('id', '').lower()
                    
                    # 水印特征检测
                    is_watermark = (
                        any(keyword in text_content for keyword in self.watermark_keywords) or
                        any(keyword in shape_id for keyword in self.watermark_keywords) or
                        ('position:absolute' in style and 'rotation:' in style and 'center' in style)
                    )
                    
                    if is_watermark:
                        watermarks['text'].append({
                            'element': shape,
                            'text': text_content,
                            'id': shape_id,
                            'style': style
                        })
            
            # 检测图片水印
            for drawing in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                anchors = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
                for anchor in anchors:
                    behind_doc = anchor.get('behindDoc', '0') == '1'
                    
                    # 检查是否居中且在文档后面
                    pos_h_center = anchor.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align')
                    pos_v_center = anchor.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align')
                    
                    h_centered = any(elem.text == 'center' for elem in pos_h_center)
                    v_centered = any(elem.text == 'center' for elem in pos_v_center)
                    
                    if behind_doc and h_centered and v_centered:
                        watermarks['image'].append({
                            'element': drawing,
                            'behind_doc': behind_doc,
                            'centered': True
                        })
            
            # 检测背景水印
            for bg in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}background'):
                vml_bg = bg.findall('.//{urn:schemas-microsoft-com:vml}background')
                if vml_bg:
                    watermarks['background'].append({
                        'element': bg
                    })
                    
        except Exception as e:
            print(f"        ⚠️  水印分析失败: {e}")
        
        return watermarks
    
    def _remove_watermarks_from_header_footer(self, header_footer, location=""):
        """从页眉页脚中移除水印 - 增强版"""
        removed_count = 0
        try:
            element = header_footer._element
            
            # 1. 检测并移除VML文本水印 (最常见的水印形式)
            for shape in element.findall('.//{urn:schemas-microsoft-com:vml}shape'):
                # 检查是否是水印形状
                textpaths = shape.findall('.//{urn:schemas-microsoft-com:vml}textpath')
                shape_id = shape.get('id', '').lower()
                style = shape.get('style', '').lower()
                
                is_watermark = False
                watermark_text = ""
                
                if textpaths:
                    watermark_text = textpaths[0].get('string', '').lower()
                    # 通过文本内容识别
                    is_watermark = any(keyword in watermark_text for keyword in self.watermark_keywords)
                    
                # 通过ID识别水印
                if not is_watermark:
                    is_watermark = any(keyword in shape_id for keyword in self.watermark_keywords)
                
                # 通过样式特征识别（绝对定位+旋转+居中）
                if not is_watermark:
                    watermark_style_features = [
                        'position:absolute', 'rotation:', 'center', 'z-index:-'
                    ]
                    feature_count = sum(1 for feature in watermark_style_features if feature in style)
                    is_watermark = feature_count >= 2  # 降低阈值，更容易检测
                
                # 特殊检测：Word标准图片水印 (如 WordPictureWatermark)
                if not is_watermark and 'picture' in shape_id and 'water' in shape_id:
                    is_watermark = True
                    watermark_text = f"Word图片水印(ID:{shape.get('id', '')})"
                
                if is_watermark:
                    shape.getparent().remove(shape)
                    removed_count += 1
                    print(f"          🗑️  移除VML文本水印 ({location}): {watermark_text[:20]}...")
            
            # 2. 移除图片水印 (w:drawing + wp:anchor)
            for drawing in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                anchors = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
                for anchor in anchors:
                    behind_doc = anchor.get('behindDoc', '0') == '1'
                    
                    # 检查定位方式
                    pos_h_center = anchor.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align')
                    pos_v_center = anchor.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align')
                    
                    # 检查图片名称
                    pic_cNvPr = anchor.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPr')
                    pic_names = [elem.get('name', '') for elem in pic_cNvPr]
                    pic_name = pic_names[0].lower() if pic_names else ""
                    
                    # 检查定位
                    h_centered = any(elem.text == 'center' for elem in pos_h_center)
                    v_centered = any(elem.text == 'center' for elem in pos_v_center)
                    
                    # 水印特征：在文档后面 + 居中 + (可能包含水印关键词)
                    is_watermark = (
                        behind_doc and 
                        h_centered and v_centered and
                        (any(keyword in pic_name for keyword in self.watermark_keywords) or 
                         'watermark' in pic_name or len(pic_name) == 0)
                    )
                    
                    if is_watermark:
                        drawing.getparent().remove(drawing)
                        removed_count += 1
                        print(f"          🗑️  移除图片水印 ({location}): {pic_name}")
                        break
            
            # 3. 移除Word Art对象
            for wordart in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object'):
                wordart.getparent().remove(wordart)
                removed_count += 1
                print(f"          🗑️  移除Word Art对象 ({location})")
            
            # 4. 移除包含VML图片的段落 (w:pict)
            for pict in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
                # 检查是否包含水印相关的VML
                vml_shapes = pict.findall('.//{urn:schemas-microsoft-com:vml}shape')
                if vml_shapes:
                    pict.getparent().remove(pict)
                    removed_count += 1
                    print(f"          🗑️  移除VML图片段落 ({location})")
            
            # 5. 移除文本水印段落
            for para in header_footer.paragraphs:
                if para.text:
                    text_lower = para.text.lower()
                    if any(keyword in text_lower for keyword in self.watermark_keywords):
                        original_text = para.text[:20]
                        para.clear()
                        removed_count += 1
                        print(f"          🗑️  移除文本水印段落 ({location}): {original_text}...")
                        
        except Exception as e:
            print(f"        ⚠️  页眉页脚水印清理失败 ({location}): {e}")
        
        return removed_count
    
    def _remove_watermarks_from_body(self):
        """从正文中移除水印元素 - 增强版"""
        removed_count = 0
        try:
            body = self.document.element.body
            if body is not None:
                # 1. 移除正文中的VML文本水印
                for shape in body.findall('.//{urn:schemas-microsoft-com:vml}shape'):
                    textpaths = shape.findall('.//{urn:schemas-microsoft-com:vml}textpath')
                    shape_id = shape.get('id', '').lower()
                    style = shape.get('style', '').lower()
                    
                    is_watermark = False
                    watermark_text = ""
                    
                    if textpaths:
                        watermark_text = textpaths[0].get('string', '').lower()
                        is_watermark = any(keyword in watermark_text for keyword in self.watermark_keywords)
                    
                    # 通过ID和样式特征识别
                    if not is_watermark:
                        is_watermark = (
                            any(keyword in shape_id for keyword in self.watermark_keywords) or
                            ('position:absolute' in style and 'z-index:-' in style)
                        )
                    
                    if is_watermark:
                        shape.getparent().remove(shape)
                        removed_count += 1
                        print(f"          🗑️  移除正文VML水印: {watermark_text[:20]}...")
                
                # 2. 移除正文中的图片水印
                for drawing in body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    anchors = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
                    for anchor in anchors:
                        behind_doc = anchor.get('behindDoc', '0') == '1'
                        
                        # 检查是否居中定位
                        pos_h_center = anchor.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align')
                        pos_v_center = anchor.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align')
                        
                        h_centered = any(elem.text == 'center' for elem in pos_h_center)
                        v_centered = any(elem.text == 'center' for elem in pos_v_center)
                        
                        if behind_doc and h_centered and v_centered:
                            drawing.getparent().remove(drawing)
                            removed_count += 1
                            print(f"          🗑️  移除正文图片水印")
                            break
                        
        except Exception as e:
            print(f"        ⚠️  正文水印清理失败: {e}")
        
        return removed_count
    
    def _remove_background_watermarks(self):
        """移除背景水印"""
        removed_count = 0
        try:
            body = self.document.element.body
            if body is not None:
                # 移除文档级背景水印
                for bg in body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}background'):
                    # 检查是否包含VML背景
                    vml_bg = bg.findall('.//{urn:schemas-microsoft-com:vml}background')
                    if vml_bg:
                        bg.getparent().remove(bg)
                        removed_count += 1
                        print(f"          🗑️  移除背景水印")
                        
        except Exception as e:
            print(f"        ⚠️  背景水印清理失败: {e}")
        
        return removed_count


def setup_a3_landscape_page(document):
    """设置A3横向页面"""
    print("    📄 设置A3横向页面...")
    try:
        for section in document.sections:
            # 设置A3横向尺寸
            section.page_width = Inches(A3_WIDTH_LANDSCAPE)
            section.page_height = Inches(A3_HEIGHT_LANDSCAPE)
            
            # 设置合理的页边距
            section.left_margin = Inches(1.0)   # 2.54cm
            section.right_margin = Inches(1.0)  # 2.54cm
            section.top_margin = Inches(1.0)    # 2.54cm
            section.bottom_margin = Inches(1.0) # 2.54cm
            
        print(f"        ✅ 页面尺寸: {A3_WIDTH_LANDSCAPE:.2f}\" × {A3_HEIGHT_LANDSCAPE:.2f}\" (A3横向)")
        return True
    except Exception as e:
        print(f"        ❌ A3页面设置失败: {e}")
        return False


# --- 增强版Excel定位器 ---

class EnhancedExcelLocator:
    """增强版Excel对象定位器

    功能说明:
        精准识别并定位DOCX文档中的嵌入式Excel对象，支持多种Excel嵌入格式。

    核心能力:
        1. 文档结构分析: 统计段落、表格、关系ID、嵌入文件等信息
        2. Excel对象识别: 基于VML+OLE+ProgID多重特征的准确识别
        3. 原表格保护: 检测段落是否在现有表格中，避免误处理
        4. 关系ID匹配: 从document.xml.rels提取Excel关系ID
        5. 说明文字检测: 识别并标记需要清理的占位符文本

    Excel对象识别条件（满足以下组合之一）:
        - ProgID="Excel.Sheet.12" + Excel关系ID
        - VML shape (ole="t") + (ProgID 或 关系ID)
        - OLE对象 + (ProgID 或 关系ID)
        - 且不在现有Word表格中

    支持的Excel格式:
        - .xlsx (Office Open XML)
        - .xlsm (启用宏的Excel)
        - .bin (OLE复合文档)
        - OLE嵌入对象 (复合文档流)

    Args:
        docx_path: DOCX文件的完整路径

    Attributes:
        docx_path: 文档路径
        namespaces: XML命名空间映射字典
    """

    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'pkg': 'http://schemas.openxmlformats.org/package/2006/relationships',
            'o': 'http://schemas.microsoft.com/office/office',
            'v': 'urn:schemas-microsoft-com:vml',
        }
    
    def analyze_document_structure(self):
        """分析文档结构"""
        print("🔍 分析文档结构...")
        
        doc = Document(self.docx_path)
        
        # 基础统计
        original_tables = len(doc.tables)
        total_paragraphs = len(doc.paragraphs)
        
        # 获取Excel相关信息
        excel_rel_ids = self._get_excel_relationships()
        embedded_files = self._get_embedded_files()
        
        print(f"    📊 文档状态:")
        print(f"        段落总数: {total_paragraphs}")
        print(f"        原始表格数: {original_tables}")
        print(f"        Excel关系ID: {excel_rel_ids}")
        print(f"        嵌入文件数: {len(embedded_files)}")
        
        return {
            'original_tables': original_tables,
            'total_paragraphs': total_paragraphs,
            'excel_rel_ids': excel_rel_ids,
            'embedded_files': embedded_files
        }
    
    def find_excel_objects_enhanced(self):
        """增强版Excel对象查找"""
        print("🎯 增强版Excel对象定位...")
        
        doc = Document(self.docx_path)
        excel_objects = []
        caption_paragraphs = []
        
        # 获取Excel关系ID
        excel_rel_ids = self._get_excel_relationships()
        
        if not excel_rel_ids:
            print("    ⚠️  未发现Excel关系，跳过处理")
            return [], []
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_xml = ET.tostring(paragraph._p, encoding='unicode')
            para_text = paragraph.text.strip()
            
            # ============ 严格的Excel对象检测 ============
            #
            # 检测策略说明：
            # 1. VML Shape检测: 查找带ole属性的VML shape元素（<v:shape ole="t">）
            #    - 支持默认命名空间和ns2命名空间前缀
            #    - ole="t" 属性标识这是一个OLE嵌入对象
            #
            # 2. OLE对象检测: 查找OLE对象元素（<o:OLEObject>）
            #    - 支持默认命名空间和ns3命名空间前缀
            #    - 这是Office特有的嵌入对象标记
            #
            # 3. ProgID检测: 查找Excel的程序标识符（ProgID="Excel.Sheet.12"）
            #    - Excel.Sheet.12 是Excel 2007+的标准ProgID
            #    - 这是识别Excel对象最可靠的特征之一
            #
            # 4. 关系ID检测: 检查段落XML是否引用了Excel相关的关系ID
            #    - 关系ID从document.xml.rels文件中提取
            #    - 关联到word/embeddings/目录下的Excel文件

            has_vml_shape = ('<v:shape' in para_xml or '<ns2:shape' in para_xml) and ('ole="t"' in para_xml)
            has_ole_object = '<o:OLEObject' in para_xml or '<ns3:OLEObject' in para_xml
            has_excel_progid = 'ProgID="Excel.Sheet.12"' in para_xml
            has_excel_relation = any(rel_id in para_xml for rel_id in excel_rel_ids)

            # ============ 安全检查：确保不在现有表格中 ============
            #
            # 原表格保护机制：
            # - 遍历段落的父节点链，检查是否在<w:tbl>表格结构中
            # - 如果段落在现有表格中，说明这是用户手动创建的表格内容
            # - 这样可以避免误处理用户表格中恰好包含的嵌入对象
            is_in_table = self._is_paragraph_in_table(paragraph, doc)

            # ============ Excel对象判断 - 放宽检测条件 ============
            #
            # 识别规则（满足以下任一组合即可）：
            #
            # 规则1: ProgID + 关系ID 组合
            #   - 适用于标准的Excel嵌入格式
            #   - 最可靠的识别方式
            #
            # 规则2: (VML Shape 或 OLE对象) + (ProgID 或 关系ID)
            #   - 适用于各种变体的Excel嵌入格式
            #   - 兼容不同版本Word创建的嵌入对象
            #
            # 排除条件: 不在现有表格中
            #   - 确保不会误处理用户手动创建的表格
            is_excel_object = (
                (has_excel_progid and has_excel_relation) or  # ProgID + Relation 组合
                ((has_vml_shape or has_ole_object) and (has_excel_progid or has_excel_relation))  # 原有逻辑
            ) and not is_in_table
            
            if is_excel_object:
                excel_objects.append({
                    'index': para_idx,
                    'paragraph': paragraph,
                    'has_vml': has_vml_shape,
                    'has_ole': has_ole_object,
                    'has_progid': has_excel_progid,
                    'has_relation': has_excel_relation,
                    'in_table': is_in_table
                })
                
                print(f"    ✅ Excel对象: 段落 {para_idx + 1}")
                print(f"        特征: VML={has_vml_shape}, OLE={has_ole_object}, ProgID={has_excel_progid}")
            
            # 说明文字检测            
            if (any(keyword in para_text for keyword in CAPTION_KEYWORDS) and
                not is_in_table and
                len(para_text) < 100):
                
                caption_paragraphs.append({
                    'index': para_idx,
                    'paragraph': paragraph,
                    'text': para_text,
                    'exact_match': para_text in CAPTION_KEYWORDS
                })
                
                print(f"    🧹 说明文字: 段落 {para_idx + 1} - {para_text}")
        
        return excel_objects, caption_paragraphs
    
    def _is_paragraph_in_table(self, paragraph, doc):
        """检查段落是否在现有表格中"""
        try:
            p_element = paragraph._p
            parent = p_element.getparent()
            
            while parent is not None:
                tag = parent.tag.split('}')[-1] if '}' in parent.tag else parent.tag
                if tag in ['tc', 'tr', 'tbl']:
                    return True
                parent = parent.getparent()
            
            return False
        except Exception:
            return False

    def _get_excel_relationships(self):
        """获取Excel关系ID"""
        excel_rel_ids = []
        
        try:
            with zipfile.ZipFile(self.docx_path, 'r') as zf:
                if 'word/_rels/document.xml.rels' in zf.namelist():
                    content = zf.read('word/_rels/document.xml.rels').decode('utf-8')
                    root = ET.fromstring(content)
                    
                    for rel in root.findall('.//pkg:Relationship', self.namespaces):
                        rel_id = rel.get('Id')
                        rel_target = rel.get('Target', '')
                        rel_type = rel.get('Type', '')
                        
                        is_excel = (
                            ('embeddings/' in rel_target and 
                             (rel_target.endswith('.xlsx') or rel_target.endswith('.xlsm') or rel_target.endswith('.bin'))) or
                            'oleObject' in rel_type
                        )
                        
                        if is_excel:
                            excel_rel_ids.append(rel_id)
        
        except Exception as e:
            print(f"    ❌ 关系分析失败: {e}")
        
        return excel_rel_ids
    
    def _get_embedded_files(self):
        """获取嵌入文件"""
        try:
            with zipfile.ZipFile(self.docx_path, 'r') as zf:
                return [f for f in zf.namelist() if f.startswith('word/embeddings/') and not f.endswith('/')]
        except Exception:
            return []


# --- Excel数据提取和表格创建 ---

def extract_embedded_excel_enhanced(docx_path):
    """增强版Excel提取"""
    print("📂 Excel数据提取...")
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            all_files = [f for f in zf.namelist() if f.startswith('word/embeddings/') and not f.endswith('/')]
            
            extracted_excels = []
            
            for file_path in all_files:
                file_data = zf.read(file_path)
                file_name = os.path.basename(file_path)
                
                # Excel文件检测
                is_excel = (
                    file_name.endswith(('.xlsx', '.xlsm')) or
                    file_data.startswith(b'\xd0\xcf\x11\xe0') or  # OLE
                    (file_data.startswith(b'PK\x03\x04') and b'xl/' in file_data[:1000]) or  # ZIP Excel
                    b'Excel' in file_data[:1000]
                )
                
                if is_excel:
                    extracted_excels.append((io.BytesIO(file_data), file_name))
                    print(f"    ✅ Excel提取: {file_name}")
            
            return extracted_excels
            
    except Exception as e:
        print(f"    ❌ 提取失败: {e}")
        return []


def get_cell_styles(cell):
    """从openpyxl单元格中提取样式信息，并打包成字典（带颜色处理）。"""
    styles = {}
    
    # 增强颜色处理逻辑
    font_color_hex = None
    if cell.font.color and cell.font.color.rgb and isinstance(cell.font.color.rgb, str):
        font_color_hex = cell.font.color.rgb[-6:]

    fill_color_hex = None
    if cell.fill.fill_type == 'solid' and cell.fill.fgColor.rgb and isinstance(cell.fill.fgColor.rgb, str):
        fill_color_hex = cell.fill.fgColor.rgb[-6:]

    styles['font'] = {
        'name': cell.font.name or 'Arial', 'size': cell.font.sz or 11, 'bold': cell.font.b or False, 
        'italic': cell.font.i or False, 'underline': cell.font.u or 'none', 'strikethrough': cell.font.strike or False,
        'color': font_color_hex  # 存储处理过的、干净的6位十六进制颜色值
    }
    styles['fill'] = {'color': fill_color_hex}
    styles['alignment'] = {'horizontal': cell.alignment.horizontal, 'vertical': cell.alignment.vertical}
    styles['number_format'] = cell.number_format
    
    return styles

def format_value(value, number_format):
    """根据Excel的number_format格式化值。这是一个简化版本，主要处理通用、数字和小数、百分比格式。"""
    if value is None:
        return ''
        
    # 如果不是数字，直接返回字符串
    if not isinstance(value, (int, float)):
        return str(value)

    # 'General' 或 None 或 Text 格式
    if number_format is None or number_format in ['General', '@', 'Text']:
        return str(value)

    try:
        # 百分比格式 (e.g., '0.00%')
        if '%' in number_format:
            decimals = 0
            if '.' in number_format:
                # 提取小数点后的位数
                decimals_part = number_format.split('.')[1]
                decimals = len(decimals_part.replace('%', ''))
            return f"{(value * 100):.{decimals}f}%"
        
        # 小数格式 (e.g., '0.00')
        # 这是一个简化的检查，仅适用于 '0.0', '0.00' 等
        if '.' in number_format and '0' in number_format and not number_format.startswith('['):
            parts = number_format.split('.')
            if len(parts) == 2 and all(c == '0' for c in parts[1]):
                decimals = len(parts[1])
                return f"{value:.{decimals}f}"
        
        # 整数格式 (e.g., '0')
        if number_format == '0':
            return str(int(round(value)))

    except Exception:
        # 如果格式化出错，返回原始值的字符串形式
        return str(value)

    # 如果没有匹配的特定格式规则，也返回原始值的字符串形式
    return str(value)

def get_excel_data_enhanced(excel_bytes, file_name):
    """增强版Excel解析，支持完整样式和格式"""
    excel_bytes.seek(0)
    
    try:
        workbook = openpyxl.load_workbook(excel_bytes, data_only=True)
        sheet = workbook.active
        table_data = []
        
        for row in sheet.iter_rows():
            row_data = []
            for cell in row:
                row_data.append({
                    'value': cell.value,
                    'styles': get_cell_styles(cell)
                })
            table_data.append(row_data)
        
        merged_ranges = [item.coord for item in sheet.merged_cells.ranges]
        print(f"        ✅ 解析成功: {sheet.max_row}行 x {sheet.max_column}列")
        return table_data, merged_ranges, sheet.max_row, sheet.max_column
        
    except Exception as e:
        print(f"        ⚠️  解析失败: {e}，使用示例数据")
    
    # 示例数据
    sample_data = [
        [
            {'value': 'Excel数据已提取', 'styles': {'font': {'name': 'Arial', 'size': 12, 'bold': True, 'color': None}, 'fill': {'color': None}, 'alignment': {'horizontal': None, 'vertical': None}, 'number_format': 'General'}},
            {'value': '(原OLE格式)', 'styles': {'font': {'name': 'Arial', 'size': 10, 'bold': False, 'color': None}, 'fill': {'color': None}, 'alignment': {'horizontal': None, 'vertical': None}, 'number_format': 'General'}}
        ]
    ]
    
    return sample_data, [], 1, 2


def optimize_table_dimensions(table, data, a3_landscape=False):
    """自适应优化表格的列宽和行高
    
    Args:
        table: Word表格对象
        data: 表格数据
        a3_landscape: 是否为A3横向布局
    """
    if not data or not data[0]:
        return
        
    # 计算每列的最大内容长度
    col_widths = []
    max_cols = max(len(row) for row in data) if data else 0
    
    for col_idx in range(max_cols):
        max_length = 0
        for row_data in data:
            if row_data and col_idx < len(row_data):
                cell_info = row_data[col_idx]
                if cell_info and 'styles' in cell_info:
                    number_format = cell_info['styles'].get('number_format', 'General')
                    display_value = format_value(cell_info.get('value'), number_format)
                    content_length = len(str(display_value)) if display_value else 0
                    max_length = max(max_length, content_length)
        col_widths.append(max_length)
    
    # 设置列宽（基于内容长度和页面布局）
    if col_widths:
        # A3横向布局时使用更大的列宽范围
        if a3_landscape:
            min_width, max_width = 2.0, 8.0  # A3横向：2.0-8.0cm范围
            width_factor = 0.4  # 更宽的列宽因子
        else:
            min_width, max_width = 1.5, 6.0  # 标准：1.5-6.0cm范围
            width_factor = 0.3
        
        for col_idx, max_length in enumerate(col_widths):
            # 基础宽度 + 内容长度因子，限制最小和最大宽度
            width_cm = min(max(min_width, max_length * width_factor), max_width)
            
            # 设置该列的所有单元格宽度
            for row in table.rows:
                if col_idx < len(row.cells):
                    row.cells[col_idx].width = Cm(width_cm)
    
    # 设置行高（基于内容和字体大小）
    for row_idx, row_data in enumerate(data):
        if row_idx < len(table.rows) and row_data:
            # 检查该行是否有多行文本或较大字体
            max_font_size = 12  # 默认字体大小
            has_multiline = False
            
            for cell_info in row_data:
                if cell_info and 'styles' in cell_info:
                    # 检查字体大小
                    font_size = cell_info['styles'].get('font', {}).get('size', 12)
                    if font_size:
                        max_font_size = max(max_font_size, font_size)
                    
                    # 检查是否有换行符
                    number_format = cell_info['styles'].get('number_format', 'General')
                    display_value = format_value(cell_info.get('value'), number_format)
                    if display_value and '\n' in str(display_value):
                        has_multiline = True
            
            # 基于字体大小和内容设置行高
            base_height = max_font_size * 0.05  # 转换为cm
            if has_multiline:
                base_height *= 2  # 多行文本增加高度
            
            # 设置最小行高（不小于0.6cm）
            row_height = max(0.6, base_height)
            try:
                table.rows[row_idx].height = Cm(row_height)
            except Exception:
                pass  # 忽略行高设置失败

def create_table_enhanced(document, data, merged_ranges, rows, cols, a3_landscape=False):
    """增强版表格创建，支持完整样式和自适应尺寸
    
    Args:
        document: Word文档对象
        data: 表格数据
        merged_ranges: 合并单元格范围
        rows: 行数
        cols: 列数
        a3_landscape: 是否为A3横向布局
    """
    table = document.add_table(rows=rows, cols=cols)
    
    try:
        table.style = 'Table Grid'
    except KeyError:
        print("  - 警告: 文档中不存在 'Table Grid' 样式。将使用默认表格样式。")
        pass
    
    # 填充数据和样式
    for r_idx, row_data in enumerate(data):
        for c_idx, cell_info in enumerate(row_data):
            if r_idx < len(table.rows) and c_idx < len(table.rows[r_idx].cells):
                try:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    # 格式化单元格值
                    number_format = cell_info['styles'].get('number_format', 'General')
                    display_value = format_value(cell_info['value'], number_format)
                    run = p.add_run(display_value)

                    font_styles = cell_info['styles']['font']
                    font = run.font
                    font.name = font_styles.get('name', 'Arial')
                    if font_styles.get('size'): 
                        font.size = Pt(font_styles['size'])
                    font.bold = font_styles.get('bold', False)
                    font.italic = font_styles.get('italic', False)
                    if font_styles.get('underline') and font_styles['underline'] != 'none': 
                        font.underline = True
                    font.strike = font_styles.get('strikethrough', False)
                    
                    # 使用处理过的干净颜色值
                    if font_styles.get('color'):
                        font.color.rgb = RGBColor.from_string(font_styles['color'])

                    fill_color = cell_info['styles']['fill'].get('color')
                    if fill_color:
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), fill_color)
                        cell._tc.get_or_add_tcPr().append(shd)

                    h_align = cell_info['styles']['alignment'].get('horizontal')
                    v_align = cell_info['styles']['alignment'].get('vertical')
                    alignment_map = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
                    p.alignment = alignment_map.get(h_align, WD_ALIGN_PARAGRAPH.LEFT)
                    valign_map = {'center': WD_ALIGN_VERTICAL.CENTER, 'bottom': WD_ALIGN_VERTICAL.BOTTOM}
                    cell.vertical_alignment = valign_map.get(v_align, WD_ALIGN_VERTICAL.TOP)
                    
                except Exception as e:
                    print(f"  - 警告: 单元格样式设置失败 [{r_idx}, {c_idx}]: {e}")
                    cell.text = str(cell_info.get('value', ''))
    
    # 处理合并单元格
    for merged_range in merged_ranges:
        try:
            min_col, min_row, max_col, max_row = openpyxl.utils.cell.range_boundaries(merged_range)
            table.cell(min_row - 1, min_col - 1).merge(table.cell(max_row - 1, max_col - 1))
        except IndexError:
            print(f"警告：处理合并单元格 '{merged_range}' 时出错，已跳过。")
        except Exception as e:
            print(f"警告：合并单元格处理失败: {e}")
    
    # 自适应列宽和行高优化
    optimize_table_dimensions(table, data, a3_landscape=a3_landscape)
    
    return table


# --- 图片渲染功能 ---

def setup_chinese_font():
    """设置中文字体支持"""
    fonts_to_try = [
        'Heiti TC', 'PingFang SC', 'STHeiti',  # macOS
        'SimHei', 'Microsoft YaHei', 'SimSun',  # Windows  
        'WenQuanYi Micro Hei', 'Noto Sans CJK SC', 'DejaVu Sans'  # Linux/fallback
    ]
    
    for font_name in fonts_to_try:
        try:
            font_path = font_manager.findfont(font_manager.FontProperties(family=font_name))
            if os.path.exists(font_path):
                plt.rcParams['font.sans-serif'] = [font_name]
                return font_name
        except Exception:
            continue
    
    # 最终fallback
    plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
    return None

def excel_to_image(excel_bytes, output_image_path, a3_landscape=False):
    """使用matplotlib将Excel字节流转换为图片文件，支持自适应布局。
    
    Args:
        excel_bytes: Excel字节流
        output_image_path: 输出图片路径
        a3_landscape: 是否使用A3横向布局优化
    """
    excel_bytes.seek(0)
    table_data, _, max_rows, max_cols = get_excel_data_enhanced(excel_bytes, "")
    if not table_data: 
        return False

    # 设置中文字体支持
    font_used = setup_chinese_font()
    if not font_used:
        print("  - 警告: 未找到合适的中文字体，中文可能显示为方框。")

    # 智能提取和格式化单元格值
    vals = []
    max_font_sizes = []
    
    for row_data in table_data:
        row_vals = []
        row_font_size = 10  # 默认字体大小
        
        for cell_info in row_data:
            # 使用格式化函数处理数值
            number_format = cell_info['styles'].get('number_format', 'General')
            display_value = format_value(cell_info['value'], number_format)
            row_vals.append(str(display_value) if display_value is not None else '')
            
            # 记录最大字体大小
            font_size = cell_info['styles']['font'].get('size', 10)
            if font_size:
                row_font_size = max(row_font_size, font_size)
        
        vals.append(row_vals)
        max_font_sizes.append(row_font_size)
    
    # 智能计算图片尺寸
    # 计算每列的实际显示宽度（考虑字体大小）
    col_widths = []
    for col_idx in range(max_cols):
        max_width = 0
        for row_idx, row_vals in enumerate(vals):
            if col_idx < len(row_vals):
                content_length = len(row_vals[col_idx])
                font_size = max_font_sizes[row_idx]
                # 字体大小影响实际显示宽度
                adjusted_width = content_length * (font_size / 10.0)
                max_width = max(max_width, adjusted_width)
        col_widths.append(max_width)
    
    # 动态计算图片尺寸
    total_content_width = sum(col_widths)
    avg_font_size = sum(max_font_sizes) / len(max_font_sizes) if max_font_sizes else 10
    
    if a3_landscape:
        # A3横向布局优化：使用A3横向比例 (420:297 ≈ 1.41:1)
        base_width = A3_WIDTH_LANDSCAPE * 0.8  # 基于A3宽度，留出边距
        base_height = A3_HEIGHT_LANDSCAPE * 0.8  # 基于A3高度，留出边距
        
        # 保持A3比例，根据内容调整
        content_factor = min(2.0, max(0.5, total_content_width * 0.02))
        fig_width = base_width * content_factor
        fig_height = base_height * content_factor
        
        print(f"        📐 A3横向布局: {fig_width:.1f}\" × {fig_height:.1f}\"")
    else:
        # 原有的动态尺寸计算
        fig_width = max(8, min(20, total_content_width * 0.12 + 2))  # 8-20英寸范围
        fig_height = max(4, min(15, max_rows * (avg_font_size / 10) * 0.5 + 1))  # 4-15英寸范围
    
    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.axis('off')  # 不显示坐标轴

    # 创建自适应表格
    table = ax.table(cellText=vals, loc='center', cellLoc='center')
    table.auto_set_font_size(False)
    
    # 设置基础字体大小（基于平均字体大小）
    base_font_size = max(8, min(14, avg_font_size * 0.8))
    table.set_fontsize(base_font_size)
    
    # 设置自适应列宽和行高
    # 计算相对列宽（归一化）
    total_width = sum(col_widths) if sum(col_widths) > 0 else 1
    relative_widths = [w / total_width for w in col_widths]
    
    for r in range(max_rows):
        for c in range(max_cols):
            if (r < len(table_data) and c < len(table_data[r]) and 
                table_data[r] and table_data[r][c] and 
                'styles' in table_data[r][c]):
                cell_info = table_data[r][c]
                table_cell = table[r, c]
                styles = cell_info.get('styles', {})
                
                # 设置单元格样式
                # 背景色处理
                fill_style = styles.get('fill', {})
                if fill_style and fill_style.get('color'):
                    table_cell.set_facecolor(f"#{fill_style['color']}")
                
                # 字体颜色
                font_style = styles.get('font', {})
                if font_style and font_style.get('color'):
                    table_cell.get_text().set_color(f"#{font_style['color']}")

                # 字体样式
                if font_style.get('bold'):
                    table_cell.get_text().set_weight('bold')
                if font_style.get('italic'):
                    table_cell.get_text().set_style('italic')
                
                # 设置单元格尺寸
                # 动态列宽（基于内容比例）
                cell_width = max(0.08, relative_widths[c] * 0.8)  # 最小宽度0.08
                table_cell.set_width(cell_width)
                
                # 动态行高（基于字体大小和内容）
                font_size = font_style.get('size', 10) if font_style else 10
                content = vals[r][c] if r < len(vals) and c < len(vals[r]) else ''
                has_multiline = '\n' in content
                
                row_height = max(0.06, (font_size / 10) * 0.08)  # 基础行高
                if has_multiline:
                    row_height *= 1.8  # 多行内容增加高度
                
                table_cell.set_height(row_height)
                
                # 文本对齐
                alignment_style = styles.get('alignment', {})
                h_align = alignment_style.get('horizontal', 'left') if alignment_style else 'left'
                alignment_map = {'center': 'center', 'right': 'right', 'justify': 'center'}
                table_cell.get_text().set_horizontalalignment(alignment_map.get(h_align, 'left'))

    # 高质量保存，确保背景处理
    save_kwargs = {
        'bbox_inches': 'tight', 
        'dpi': 300, 
        'edgecolor': 'none', 
        'pad_inches': 0.1
    }
    
    # 设置背景色
    save_kwargs['facecolor'] = 'white'
    
    plt.savefig(output_image_path, **save_kwargs)
    plt.close(fig)  # 关闭图形，释放内存
    
    print(f"        ✅ 图片已保存: {output_image_path}")
    
    return True


# --- 内置扩展: DOCX表格提取 ---

def insert_paragraph_before_table(table, text):
    """在表格前插入段落标记"""
    tbl_element = table._element
    parent = tbl_element.getparent()
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    parent.insert(parent.index(tbl_element), p)


def extract_table_data(table):
    """提取表格数据为二维列表"""
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip().replace('\n', ' '))
        data.append(row_data)
    return data


def save_to_txt(tables_data, output_path):
    """保存所有表格为 TXT（CSV）"""
    with open(output_path, 'w', encoding='utf-8') as f:
        for idx, data in tables_data:
            f.write(f"【表{idx}】\n")
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerows(data)
            f.write(output.getvalue())
            f.write("\n" + "=" * 50 + "\n\n")
    print(f"  ✓ 已导出 TXT: {output_path.name}")


def save_to_xlsx(tables_data, output_path):
    """保存所有表格为 XLSX"""
    wb = openpyxl.Workbook()
    if 'Sheet' in wb.sheetnames:
        del wb['Sheet']

    ws_all = wb.create_sheet("所有表格")
    current_row = 1

    for idx, data in tables_data:
        cell = ws_all.cell(row=current_row, column=1, value=f"【表{idx}】")
        cell.font = openpyxl.styles.Font(bold=True, size=12)
        current_row += 1

        for row_data in data:
            for col_idx, cell_value in enumerate(row_data, 1):
                ws_all.cell(row=current_row, column=col_idx, value=cell_value)
            current_row += 1

        current_row += 2

    wb.save(output_path)
    print(f"  ✓ 已导出 XLSX: {output_path.name}")


class PageMarker(Flowable):
    """PDF渲染时记录真实页码"""

    def __init__(self, key, page_tracker):
        Flowable.__init__(self)
        self.key = key
        self.page_tracker = page_tracker
        self.width = 0
        self.height = 0

    def draw(self):
        self.page_tracker[self.key] = self.canv.getPageNumber()


# ReportLab 中文字体缓存（避免重复注册）
_reportlab_font_cache = {}

# ReportLab 字体搜索路径（macOS + Windows 全覆盖）
_REPORTLAB_FONT_PATHS = [
    "/System/Library/Fonts/STHeiti Medium.ttc",   # macOS
    "/System/Library/Fonts/STHeiti Light.ttc",     # macOS
    "/System/Library/Fonts/PingFang.ttc",          # macOS
    "/Library/Fonts/Songti.ttc",                   # macOS
    "C:\\Windows\\Fonts\\simhei.ttf",              # Windows
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",                # Linux (WenQuanYi)
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",      # Linux (Noto CJK)
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",         # Linux (Noto CJK)
]


def register_reportlab_chinese_font():
    """注册中文字体供 ReportLab PDF 导出使用（表格PDF + 图片分离PDF 通用）"""
    if _reportlab_font_cache:
        return _reportlab_font_cache['name']

    font_name = "ChinesePDF"
    for path in _REPORTLAB_FONT_PATHS:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, path))
                _reportlab_font_cache['name'] = font_name
                return font_name
            except Exception:
                continue

    print("⚠️  未找到合适的中文字体，PDF 中文可能无法显示。")
    _reportlab_font_cache['name'] = "Helvetica"
    return "Helvetica"


def calculate_smart_col_widths(data, font_name, max_width=None, font_size=10):
    """智能计算列宽"""
    if not data:
        return []

    num_cols = len(data[0])
    desired_widths = [0] * num_cols
    max_single_col_width = 500

    for row in data:
        for i, cell in enumerate(row):
            if i < num_cols:
                width = 0
                for char in str(cell):
                    if '\u4e00' <= char <= '\u9fff':
                        width += font_size
                    elif char.isupper():
                        width += font_size * 0.8
                    else:
                        width += font_size * 0.65
                width += 16
                desired_widths[i] = max(desired_widths[i], min(width, max_single_col_width))

    total_desired = sum(desired_widths)
    if max_width is None or total_desired <= max_width:
        return desired_widths

    min_col_width = font_size * 4
    min_total_width = num_cols * min_col_width
    if min_total_width > max_width:
        return [max(w, min_col_width) for w in desired_widths]

    current_total = total_desired
    final_widths = []
    for w in desired_widths:
        ratio = w / current_total
        allocated = max_width * ratio
        final_widths.append(max(allocated, min_col_width))

    return final_widths


def calculate_optimal_page_size(data, font_name):
    """计算表格PDF最佳页面尺寸"""
    if not data:
        return A4

    widths = calculate_smart_col_widths(data, font_name, max_width=None)
    table_width = sum(widths) * 1.05

    max_cell_height = 0
    for row in data:
        for i, cell in enumerate(row):
            if i < len(widths):
                col_w = widths[i]
                text_len = len(str(cell))
                est_lines = (text_len * 10 * 0.8) / (col_w - 8)
                est_lines = max(1, int(est_lines) + 1)
                cell_h = est_lines * 14
                max_cell_height = max(max_cell_height, cell_h)

    required_h = max_cell_height + 4 * inch

    max_page_width_limit = 14000
    max_page_height_limit = 14000

    page_width = max(table_width + 2 * inch, 842)
    page_height = max(required_h, 595)

    page_width = min(page_width, max_page_width_limit)
    page_height = min(page_height, max_page_height_limit)

    return (page_width, page_height)


def _make_table_style(font_name):
    """创建PDF表格统一样式"""
    return TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), font_name),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
    ])


def _build_table_story_elements(idx, data, table_sizes, font_name, style_cn, style_title,
                                page_tracker=None):
    """构建单个表格的PDF story元素列表。page_tracker 非None时插入PageMarker（第一遍分析用）"""
    elements = []
    elements.append(NextPageTemplate(f'PT_{idx}'))
    if page_tracker is not None:
        elements.append(PageMarker(f'table_{idx}_start', page_tracker))
    elements.append(Paragraph(f"<a name='Table_{idx}'/>【表{idx}】", style_title))

    p_w, p_h = table_sizes[idx]
    available_width = p_w - inch
    col_widths = calculate_smart_col_widths(data, font_name, max_width=available_width)

    table_data = [[Paragraph(str(cell), style_cn) for cell in row] for row in data]
    t = LongTable(table_data, colWidths=col_widths, repeatRows=1, hAlign='LEFT', splitInRow=1)
    t.setStyle(_make_table_style(font_name))

    elements.append(t)
    elements.append(Spacer(1, 0.5 * inch))
    if page_tracker is not None:
        elements.append(PageMarker(f'table_{idx}_end', page_tracker))
    elements.append(PageBreak())
    return elements


def save_to_pdf(tables_data, output_path):
    """保存所有表格为PDF（自适应页面+真实页码目录）"""
    font_name = register_reportlab_chinese_font()

    styles = getSampleStyleSheet()
    style_cn = ParagraphStyle(
        name='ChineseStyle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=14,
        wordWrap='CJK',
        alignment=TA_LEFT,
    )

    style_title = ParagraphStyle(
        name='TitleStyle',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=14,
        leading=18,
        spaceAfter=12,
        alignment=TA_LEFT,
    )

    style_toc = ParagraphStyle(
        name='TOCStyle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=12,
        leading=16,
        spaceAfter=6,
    )

    page_templates = []
    toc_frame = Frame(0.5 * inch, 0.5 * inch, A4[0] - inch, A4[1] - inch, id='toc_frame')
    page_templates.append(PageTemplate(id='TOC', frames=[toc_frame], pagesize=A4))

    table_sizes = {}
    for idx, data in tables_data:
        if not data:
            continue

        p_w, p_h = calculate_optimal_page_size(data, font_name)
        table_sizes[idx] = (p_w, p_h)

        margin = 0.5 * inch
        frame = Frame(margin, margin, p_w - 2 * margin, p_h - 2 * margin, id=f'frame_{idx}')
        template = PageTemplate(id=f'PT_{idx}', frames=[frame], pagesize=(p_w, p_h))
        page_templates.append(template)

    doc = BaseDocTemplate(str(output_path), pageTemplates=page_templates)

    print("  🔍 第一遍：分析表格页码...")
    page_tracker = {}
    story_first = []

    story_first.append(NextPageTemplate('TOC'))
    story_first.append(Paragraph("表格目录 / Table of Contents（正在分析...）", style_title))
    story_first.append(Spacer(1, 0.5 * inch))

    for idx, data in tables_data:
        if not data:
            continue
        num_rows = len(data)
        story_first.append(
            Paragraph(f"• 【表{idx}】............ 第 999-999 页（{num_rows} 行，共 99 页）", style_toc)
        )
    story_first.append(PageBreak())

    for idx, data in tables_data:
        if not data:
            continue
        story_first.extend(_build_table_story_elements(
            idx, data, table_sizes, font_name, style_cn, style_title, page_tracker=page_tracker
        ))

    temp_output = str(output_path.with_name(output_path.stem + '_temp.pdf'))
    doc_first = BaseDocTemplate(temp_output, pageTemplates=page_templates)

    page_analysis_ok = True
    try:
        doc_first.build(story_first)
        print(f"  ✓ 页码分析完成，发现 {len(page_tracker) // 2} 个表格")
    except Exception as e:
        page_analysis_ok = False
        print(f"  ⚠️  页码分析失败，将生成不含真实页码的目录: {e}")

    print("  📝 第二遍：生成最终 PDF...")
    story = []
    story.append(NextPageTemplate('TOC'))

    table_page_ranges = []
    total_pages = 1
    if page_analysis_ok and page_tracker:
        for idx, data in tables_data:
            if not data:
                continue

            start_key = f'table_{idx}_start'
            end_key = f'table_{idx}_end'
            if start_key in page_tracker and end_key in page_tracker:
                start_page = page_tracker[start_key]
                end_page = page_tracker[end_key]
                num_rows = len(data)
                table_page_ranges.append((idx, start_page, end_page, num_rows))
        total_pages = max(page_tracker.values())
        story.append(Paragraph(f"表格目录 / Table of Contents（共 {total_pages} 页）", style_title))
    else:
        story.append(Paragraph("表格目录 / Table of Contents（无真实页码）", style_title))
    story.append(Spacer(1, 0.5 * inch))

    if page_analysis_ok and table_page_ranges:
        for idx, start_page, end_page, num_rows in table_page_ranges:
            if start_page == end_page:
                page_info = f"第 {start_page} 页（{num_rows} 行）"
            else:
                page_info = f"第 {start_page}-{end_page} 页（{num_rows} 行，共 {end_page - start_page + 1} 页）"

            story.append(
                Paragraph(f"• <a href='#Table_{idx}'>【表{idx}】</a> ............ {page_info}", style_toc)
            )
    else:
        for idx, data in tables_data:
            if not data:
                continue
            num_rows = len(data)
            story.append(
                Paragraph(f"• <a href='#Table_{idx}'>【表{idx}】</a> ............ （{num_rows} 行）", style_toc)
            )
    story.append(PageBreak())

    for idx, data in tables_data:
        if not data:
            continue
        story.extend(_build_table_story_elements(
            idx, data, table_sizes, font_name, style_cn, style_title, page_tracker=None
        ))

    try:
        doc.build(story)
        print(f"  ✓ 已导出 PDF: {output_path.name}")
    except Exception as e:
        print(f"❌ PDF 生成失败: {e}")
    finally:
        if os.path.exists(temp_output):
            try:
                os.remove(temp_output)
            except Exception:
                pass


def process_docx(docx_path):
    """处理DOCX并提取表格到 TXT/XLSX/PDF"""
    docx_path = Path(docx_path)
    if not docx_path.exists():
        print(f"❌ 文件不存在: {docx_path}")
        return False
    if docx_path.is_dir():
        print(f"❌ 输入是文件夹，不是 DOCX 文件: {docx_path}")
        return False

    already_marked_table = docx_path.stem.endswith(TABLE_MARK_SUFFIX)
    if already_marked_table:
        print(f"🏷️  检测到已标记表格文件，将直接导出（不重复插入【表N】）: {docx_path.name}")

    print(f"📄 处理文件: {docx_path.name}")
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"❌ 无法打开文档: {e}")
        return False

    tables_data = []
    tables = list(doc.tables)
    if not tables:
        print("⚠️  文档中没有找到表格")
        return None

    print(f"  📊 发现 {len(tables)} 个表格")
    for idx, table in enumerate(tables, 1):
        data = extract_table_data(table)
        tables_data.append((idx, data))
        if not already_marked_table:
            insert_paragraph_before_table(table, f"【表{idx}】")
        print(f"    处理 表{idx} ({len(data)}行)")

    if not already_marked_table:
        output_docx_path = docx_path.parent / f"{docx_path.stem}{TABLE_MARK_SUFFIX}.docx"
        doc.save(output_docx_path)
        print(f"  ✓ 已保存标记文档: {output_docx_path.name}")

    output_txt_path = docx_path.parent / f"{docx_path.stem}_表格提取.txt"
    output_xlsx_path = docx_path.parent / f"{docx_path.stem}_表格提取.xlsx"
    output_pdf_path = docx_path.parent / f"{docx_path.stem}_表格提取.pdf"

    save_to_txt(tables_data, output_txt_path)
    save_to_xlsx(tables_data, output_xlsx_path)
    save_to_pdf(tables_data, output_pdf_path)

    print("\n✅ 处理完成!")
    return True


# --- 内置扩展: DOCX图片分离 ---


def _extract_images_from_element(element, location_prefix, all_images, used_rids, active_images):
    """从 XML 元素中提取图片引用（DrawingML + VML）"""
    for drawing in element.findall('.//' + qn('w:drawing')):
        blip = drawing.find('.//' + qn('a:blip'))
        if blip is not None:
            embed_id = blip.get(qn('r:embed'))
            if embed_id and embed_id in all_images and embed_id not in used_rids:
                used_rids.add(embed_id)
                active_images.append((embed_id, location_prefix))

    for pict in element.findall('.//' + qn('w:pict')):
        for imagedata in pict.findall('.//' + qn('v:imagedata')):
            embed_id = imagedata.get(qn('r:id')) or imagedata.get(qn('o:relid'))
            if embed_id and embed_id in all_images and embed_id not in used_rids:
                used_rids.add(embed_id)
                active_images.append((embed_id, f"{location_prefix}[VML]"))


def analyze_document_images(docx_path):
    """分析文档图片：有效图片、孤儿图片、资源总表"""
    doc = Document(docx_path)
    all_images = {}

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                content_type = rel.target_part.content_type
                image_format = content_type.split('/')[-1]
                if image_format == 'jpeg':
                    image_format = 'jpg'
                all_images[rel.rId] = {
                    'data': rel.target_part.blob,
                    'format': image_format,
                    'size': len(rel.target_part.blob),
                    'target': rel.target_ref
                }
            except Exception as e:
                print(f"  ⚠️  跳过无效图片关系 {rel.rId}: {e}")

    for section in doc.sections:
        try:
            header = section.header
            if header and hasattr(header, 'part') and hasattr(header.part, 'rels'):
                for rel in header.part.rels.values():
                    if "image" in rel.reltype and rel.rId not in all_images:
                        try:
                            content_type = rel.target_part.content_type
                            image_format = content_type.split('/')[-1]
                            if image_format == 'jpeg':
                                image_format = 'jpg'
                            all_images[rel.rId] = {
                                'data': rel.target_part.blob,
                                'format': image_format,
                                'size': len(rel.target_part.blob),
                                'target': rel.target_ref
                            }
                        except Exception:
                            pass
        except Exception:
            pass

        try:
            footer = section.footer
            if footer and hasattr(footer, 'part') and hasattr(footer.part, 'rels'):
                for rel in footer.part.rels.values():
                    if "image" in rel.reltype and rel.rId not in all_images:
                        try:
                            content_type = rel.target_part.content_type
                            image_format = content_type.split('/')[-1]
                            if image_format == 'jpeg':
                                image_format = 'jpg'
                            all_images[rel.rId] = {
                                'data': rel.target_part.blob,
                                'format': image_format,
                                'size': len(rel.target_part.blob),
                                'target': rel.target_ref
                            }
                        except Exception:
                            pass
        except Exception:
            pass

    active_images = []
    used_rids = set()

    for para_idx, paragraph in enumerate(doc.paragraphs):
        text_preview = paragraph.text.strip()[:50]
        location = f"正文-段落{para_idx}: {text_preview}" if text_preview else f"正文-段落{para_idx}"
        _extract_images_from_element(paragraph._element, location, all_images, used_rids, active_images)

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    location = f"正文-表格{table_idx+1}-行{row_idx+1}-列{cell_idx+1}"
                    _extract_images_from_element(para._element, location, all_images, used_rids, active_images)

    for section_idx, section in enumerate(doc.sections):
        for header_name in ["首页页眉", "偶数页页眉", "默认页眉"]:
            try:
                header = section.header
                if header:
                    for para in header.paragraphs:
                        location = f"第{section_idx+1}节-{header_name}"
                        _extract_images_from_element(para._element, location, all_images, used_rids, active_images)
                    for table_idx, table in enumerate(header.tables):
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    location = f"第{section_idx+1}节-{header_name}-表格{table_idx+1}"
                                    _extract_images_from_element(
                                        para._element, location, all_images, used_rids, active_images
                                    )
            except Exception:
                pass

        for footer_name in ["首页页脚", "偶数页页脚", "默认页脚"]:
            try:
                footer = section.footer
                if footer:
                    for para in footer.paragraphs:
                        location = f"第{section_idx+1}节-{footer_name}"
                        _extract_images_from_element(para._element, location, all_images, used_rids, active_images)
                    for table_idx, table in enumerate(footer.tables):
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    location = f"第{section_idx+1}节-{footer_name}-表格{table_idx+1}"
                                    _extract_images_from_element(
                                        para._element, location, all_images, used_rids, active_images
                                    )
            except Exception:
                pass

    try:
        body_element = doc.element.body
        for txbx_content in body_element.findall('.//' + qn('w:txbxContent')):
            for para_idx, para_element in enumerate(txbx_content.findall('.//' + qn('w:p'))):
                location = f"文本框-段落{para_idx}"
                _extract_images_from_element(para_element, location, all_images, used_rids, active_images)
    except Exception:
        pass

    orphan_images = [rid for rid in all_images.keys() if rid not in used_rids]
    return {
        'active_images': active_images,
        'orphan_images': orphan_images,
        'all_images': all_images
    }


def _catalog_new_page(c, font_name, page_num, header_text=None):
    """目录翻页：绘制页脚、换页、设置页面，返回 (new_page_num, new_y)"""
    c.setFont(font_name, 8)
    c.setFillColor(HexColor('#999999'))
    c.drawString(60, 40, "生成工具: DOCX图片分离工具 - 内置版")
    c.drawString(IMG_A3_WIDTH - 250, 40, f"目录第 {page_num} 页")
    c.showPage()

    c.setPageSize((IMG_A3_WIDTH, IMG_A3_HEIGHT))
    page_num += 1
    y = IMG_A3_HEIGHT - 60
    if header_text:
        c.setFont(font_name, 14)
        c.setFillColor(HexColor('#1a1a1a'))
        c.drawString(60, y, header_text)
        y -= 35
        c.setFont(font_name, 10)
        c.setFillColor(HexColor('#555555'))
    return page_num, y


def create_catalog_pages(c, analysis_result, font_name):
    """创建图片目录页（A3纵向，多页）"""
    active_images = analysis_result['active_images']
    active_count = len(active_images)
    orphan_count = len(analysis_result['orphan_images'])
    total_count = len(analysis_result['all_images'])

    c.setPageSize((IMG_A3_WIDTH, IMG_A3_HEIGHT))
    c.setFont(font_name, 24)
    c.setFillColor(HexColor('#1a1a1a'))
    c.drawString(60, IMG_A3_HEIGHT - 60, "图片索引目录")

    c.setStrokeColor(HexColor('#cccccc'))
    c.setLineWidth(1)
    c.line(60, IMG_A3_HEIGHT - 75, IMG_A3_WIDTH - 60, IMG_A3_HEIGHT - 75)

    c.setFont(font_name, 14)
    c.setFillColor(HexColor('#333333'))
    y = IMG_A3_HEIGHT - 110
    c.drawString(60, y, f"文档统计:")
    y -= 30
    c.setFont(font_name, 12)
    c.drawString(80, y, f"• 总图片数: {total_count} 张")
    y -= 25
    c.drawString(80, y, f"• 有效图片: {active_count} 张 (已标记并连续编号)")
    y -= 25
    c.drawString(80, y, f"• 孤儿图片: {orphan_count} 张 (未被引用)")

    y -= 50
    c.setFont(font_name, 14)
    c.setFillColor(HexColor('#1a1a1a'))
    c.drawString(60, y, "有效图片索引:")

    y -= 35
    c.setFont(font_name, 10)
    c.setFillColor(HexColor('#555555'))

    page_num = 1
    items_per_page_first = 35
    items_per_page_rest = 50
    _cont_header = "有效图片索引 (续):"

    # 预计算目录总页数，用于准确标注图片所在PDF页码
    if active_count <= items_per_page_first:
        est_catalog_pages = 1
    else:
        est_catalog_pages = 1 + -(-max(0, active_count - items_per_page_first) // items_per_page_rest)
    # 预估孤儿区是否会导致额外翻页（孤儿区约需 200pt，检查最后一页剩余空间）
    if orphan_count > 0:
        if active_count <= items_per_page_first:
            last_page_items = active_count
            last_page_capacity = items_per_page_first
        else:
            last_page_items = (active_count - items_per_page_first) % items_per_page_rest
            if last_page_items == 0:
                last_page_items = items_per_page_rest
            last_page_capacity = items_per_page_rest
        remaining_y_est = (last_page_capacity - last_page_items) * 20
        if remaining_y_est < 230:
            est_catalog_pages += 1

    for idx, (rel_id, location) in enumerate(active_images, 1):
        needs_break = False
        if page_num == 1 and idx > items_per_page_first:
            needs_break = True
        elif page_num > 1 and (idx - items_per_page_first - 1) % items_per_page_rest == 0 and idx > items_per_page_first:
            needs_break = True

        if needs_break or y < 80:
            page_num, y = _catalog_new_page(c, font_name, page_num, header_text=_cont_header)

        pdf_page = idx + est_catalog_pages
        text = f"图{idx} → PDF第{pdf_page}页 | 位置: {location}"
        if len(text) > 110:
            text = text[:107] + "..."

        c.drawString(70, y, text)
        y -= 20

    if orphan_count > 0:
        y -= 30
        if y < 200:
            page_num, y = _catalog_new_page(c, font_name, page_num)

        c.setFont(font_name, 14)
        c.setFillColor(HexColor('#d32f2f'))
        c.drawString(60, y, f"孤儿图片 ({orphan_count} 张):")

        y -= 30
        c.setFont(font_name, 11)
        c.setFillColor(HexColor('#666666'))
        orphan_start_page = active_count + page_num + 1
        c.drawString(70, y, f"未被文档引用的图片已附在PDF末尾 (第{orphan_start_page}页起)")
        y -= 22
        c.drawString(70, y, "这些图片可能是:")
        y -= 20
        c.drawString(85, y, "• 删除后残留的图片资源")
        y -= 20
        c.drawString(85, y, "• 被其他元素（页眉/页脚/文本框）引用")
        y -= 20
        c.drawString(85, y, "• 重复导入但未使用的图片")

    c.setFont(font_name, 8)
    c.setFillColor(HexColor('#999999'))
    c.drawString(60, 40, "生成工具: DOCX图片分离工具 - 内置版")
    c.drawString(IMG_A3_WIDTH - 250, 40, f"目录第 {page_num} 页 / 共 {page_num} 页")
    c.showPage()
    return page_num


def _create_marker_run(text):
    """创建标记run元素（安全XML构建，避免注入风险）"""
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    return r


def _mark_images_in_element(element, rel_id_to_index, remove_images):
    """在XML段落元素中插入图片标记（支持DrawingML和VML）"""
    count = 0

    for run_element in element.findall('./' + qn('w:r')):
        for drawing in run_element.findall('.//' + qn('w:drawing')):
            blip = drawing.find('.//' + qn('a:blip'))
            if blip is not None:
                embed_id = blip.get(qn('r:embed'))
                if embed_id in rel_id_to_index:
                    run_index = list(element).index(run_element)
                    element.insert(run_index, _create_marker_run(f'【图{rel_id_to_index[embed_id]}】'))
                    if remove_images:
                        run_element.remove(drawing)
                    count += 1

    for run_element in element.findall('./' + qn('w:r')):
        for pict in run_element.findall('.//' + qn('w:pict')):
            for imagedata in pict.findall('.//' + qn('v:imagedata')):
                embed_id = imagedata.get(qn('r:id'))
                if embed_id in rel_id_to_index:
                    run_index = list(element).index(run_element)
                    element.insert(run_index, _create_marker_run(f'【图{rel_id_to_index[embed_id]}】'))
                    if remove_images:
                        run_element.remove(pict)
                    count += 1

    return count


def mark_images_in_docx(docx_path, output_path, analysis_result, remove_images=False):
    """在DOCX正文/表格/页眉页脚/文本框中标记图片"""
    doc = Document(docx_path)
    rel_id_to_index = {}
    for idx, (rel_id, _) in enumerate(analysis_result['active_images'], 1):
        if rel_id not in rel_id_to_index:
            rel_id_to_index[rel_id] = idx

    replaced_count = 0

    for paragraph in doc.paragraphs:
        replaced_count += _mark_images_in_element(paragraph._element, rel_id_to_index, remove_images)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replaced_count += _mark_images_in_element(paragraph._element, rel_id_to_index, remove_images)

    for section in doc.sections:
        try:
            header = section.header
            if header:
                for paragraph in header.paragraphs:
                    replaced_count += _mark_images_in_element(paragraph._element, rel_id_to_index, remove_images)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replaced_count += _mark_images_in_element(
                                    paragraph._element, rel_id_to_index, remove_images
                                )
        except Exception:
            pass

        try:
            footer = section.footer
            if footer:
                for paragraph in footer.paragraphs:
                    replaced_count += _mark_images_in_element(paragraph._element, rel_id_to_index, remove_images)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replaced_count += _mark_images_in_element(
                                    paragraph._element, rel_id_to_index, remove_images
                                )
        except Exception:
            pass

    try:
        body_element = doc.element.body
        for txbx_content in body_element.findall('.//' + qn('w:txbxContent')):
            for para_element in txbx_content.findall('.//' + qn('w:p')):
                replaced_count += _mark_images_in_element(para_element, rel_id_to_index, remove_images)
    except Exception:
        pass

    doc.save(output_path)
    return replaced_count


def calculate_page_size(image_width, image_height):
    """计算图片PDF页面尺寸"""
    img_width_pt = image_width
    img_height_pt = image_height

    if img_width_pt >= IMG_A4_WIDTH and img_height_pt >= IMG_A4_HEIGHT:
        return img_width_pt, img_height_pt, 1.0

    width_scale = IMG_A4_WIDTH / img_width_pt
    height_scale = IMG_A4_HEIGHT / img_height_pt
    scale = min(width_scale, height_scale)

    page_width = img_width_pt * scale
    page_height = img_height_pt * scale
    return page_width, page_height, scale


def optimize_image_for_pdf(img, original_format='png', quality=85):
    """根据图像特征优化输出格式（JPEG/PNG）"""
    img_buffer = io.BytesIO()

    if img.mode == 'RGBA':
        alpha = img.split()[3]
        if alpha.getextrema() == (255, 255):
            rgb_img = Image.new('RGB', img.size, (255, 255, 255))
            rgb_img.paste(img, mask=img.split()[3])
            img = rgb_img
            use_jpeg = True
        else:
            use_jpeg = False
    elif img.mode in ['RGB', 'L']:
        use_jpeg = True
    else:
        img = img.convert('RGB')
        use_jpeg = True

    if original_format == 'jpeg' and use_jpeg:
        img.save(img_buffer, format='JPEG', quality=quality, optimize=True)
        return img_buffer, 'JPEG'

    if use_jpeg and img.mode in ['RGB', 'L']:
        try:
            colors_result = img.getcolors(maxcolors=256)
            if colors_result is None or len(colors_result) > 128:
                img.save(img_buffer, format='JPEG', quality=quality, optimize=True)
                return img_buffer, 'JPEG'
            img.save(img_buffer, format='PNG', optimize=True)
            return img_buffer, 'PNG'
        except Exception:
            img.save(img_buffer, format='JPEG', quality=quality, optimize=True)
            return img_buffer, 'JPEG'

    img.save(img_buffer, format='PNG', optimize=True)
    return img_buffer, 'PNG'


def _render_image_page(c, img_info, font_name, optimize, jpeg_quality,
                       label_text, label_font_size=8, label_color='#666666'):
    """渲染单张图片到PDF页面，返回 (original_size, optimized_size) 或 None（失败时）"""
    try:
        image_data = img_info['data']
        original_size = len(image_data)

        img = Image.open(io.BytesIO(image_data))
        img_width, img_height = img.size
        page_width, page_height, scale = calculate_page_size(img_width, img_height)
        c.setPageSize((page_width, page_height))

        if optimize:
            img_buffer, final_format = optimize_image_for_pdf(
                img, original_format=img_info['format'], quality=jpeg_quality
            )
        else:
            if img.mode == 'RGBA':
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                rgb_img.paste(img, mask=img.split()[3])
                img = rgb_img
            elif img.mode not in ['RGB', 'L']:
                img = img.convert('RGB')
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            final_format = 'PNG'

        img_buffer.seek(0)
        optimized_size = len(img_buffer.getvalue())
        img_reader = ImageReader(img_buffer)

        x = (page_width - img_width * scale) / 2
        y = (page_height - img_height * scale) / 2
        c.drawImage(img_reader, x, y, width=img_width * scale, height=img_height * scale)

        c.setFont(font_name, label_font_size)
        c.setFillColor(HexColor(label_color))
        c.drawString(10, page_height - 15, label_text)
        c.showPage()

        return original_size, optimized_size, img_width, img_height, final_format
    except Exception as e:
        print(f"    ... 图片渲染异常: {e}")
        return None


def create_pdf_with_catalog(analysis_result, output_pdf_path, optimize=True, jpeg_quality=85):
    """创建图片索引PDF"""
    all_images = analysis_result['all_images']
    active_images = analysis_result['active_images']
    orphan_images = analysis_result['orphan_images']

    if not all_images:
        print("⚠️  没有图片可以生成 PDF")
        return False

    font_name = register_reportlab_chinese_font()
    c = canvas.Canvas(str(output_pdf_path), pagesize=(IMG_A3_WIDTH, IMG_A3_HEIGHT))

    print("  📑 生成目录页...")
    catalog_pages = create_catalog_pages(c, analysis_result, font_name)
    print(f"     目录共 {catalog_pages} 页")

    print(f"\n  📸 添加有效图片 ({len(active_images)} 张):")
    if optimize:
        print(f"     优化模式: JPEG质量={jpeg_quality}, 智能格式选择")

    total_original_size = 0
    total_optimized_size = 0

    for idx, (rel_id, location) in enumerate(active_images, 1):
        label = f"图{idx} | {location[:40]}"
        result = _render_image_page(c, all_images[rel_id], font_name, optimize, jpeg_quality,
                                    label_text=label)
        if result:
            original_size, optimized_size, img_width, img_height, final_format = result
            total_original_size += original_size
            total_optimized_size += optimized_size
            if optimize:
                ratio = (1 - optimized_size / original_size) * 100
                print(
                    f"    ✓ 图{idx}: {img_width}x{img_height}px | {final_format} | "
                    f"{original_size//1024}KB→{optimized_size//1024}KB ({ratio:+.0f}%) | {location[:30]}"
                )
            else:
                print(f"    ✓ 图{idx}: {img_width}x{img_height}px | {location[:50]}")
        else:
            print(f"    ❌ 图{idx} 处理失败")

    if orphan_images:
        print(f"\n  🗑️  添加孤儿图片 ({len(orphan_images)} 张):")
        for rel_id in orphan_images:
            img_info = all_images[rel_id]
            label = f"[未使用] {rel_id} | {img_info['size']//1024}KB"
            result = _render_image_page(c, img_info, font_name, optimize, jpeg_quality,
                                        label_text=label, label_font_size=10, label_color='#d32f2f')
            if result:
                original_size, optimized_size, img_width, img_height, final_format = result
                total_original_size += original_size
                total_optimized_size += optimized_size
                if optimize:
                    ratio = (1 - optimized_size / original_size) * 100
                    print(
                        f"    • {rel_id}: {img_width}x{img_height}px | {final_format} | "
                        f"{original_size//1024}KB→{optimized_size//1024}KB ({ratio:+.0f}%)"
                    )
                else:
                    print(f"    • {rel_id}: {img_width}x{img_height}px ({img_info['size']//1024}KB)")
            else:
                print(f"    ❌ {rel_id} 处理失败")

    c.save()

    if optimize and total_original_size > 0:
        compression_ratio = (1 - total_optimized_size / total_original_size) * 100
        print(f"\n📊 图片优化统计:")
        print(f"   原始总大小: {total_original_size/1024/1024:.2f} MB")
        print(f"   优化后大小: {total_optimized_size/1024/1024:.2f} MB")
        print(f"   压缩率: {compression_ratio:.1f}%")

    print(f"\n✅ PDF 生成成功: {output_pdf_path}")
    return True


def process_docx_file(docx_path, remove_images=False, output_dir=None, optimize_images=True, jpeg_quality=85):
    """处理单个DOCX的图片分离流程"""
    docx_path = Path(docx_path)
    if not docx_path.exists():
        print(f"❌ 文件不存在: {docx_path}")
        return False
    if docx_path.suffix.lower() != '.docx':
        print(f"❌ 不是 DOCX 文件: {docx_path}")
        return False

    print(f"\n📄 处理文件: {docx_path.name}")
    print("=" * 80)

    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
    else:
        output_dir = docx_path.parent

    base_name = docx_path.stem
    output_docx_path = output_dir / f"{base_name}_已标记图片.docx"
    output_pdf_path = output_dir / f"{base_name}_附图.pdf"

    try:
        print("🔍 分析文档图片...")
        analysis_result = analyze_document_images(docx_path)

        active_count = len(analysis_result['active_images'])
        orphan_count = len(analysis_result['orphan_images'])
        total_count = len(analysis_result['all_images'])

        print(f"  • 图片资源总数: {total_count}")
        print(f"  • 有效图片: {active_count} (被文档引用)")
        print(f"  • 孤儿图片: {orphan_count} (未被引用)")

        if total_count == 0:
            print("⚠️  未找到图片")
            return False

        print("\n📚 生成PDF...")
        create_pdf_with_catalog(
            analysis_result, output_pdf_path, optimize=optimize_images, jpeg_quality=jpeg_quality
        )

        print("\n🏷️  标记图片位置...")
        replaced_count = mark_images_in_docx(
            docx_path, output_docx_path, analysis_result, remove_images
        )
        print(f"  ✓ 已标记 {replaced_count} 处图片位置 (对应 {active_count} 张唯一图片，编号: 图1~图{active_count})")

        if remove_images:
            print("  ✓ 已删除原文档中的图片")

        print("\n" + "=" * 80)
        print("✅ 处理完成!")
        print(f"  📝 DOCX: {output_docx_path.name}")
        print(f"  📄 PDF:  {output_pdf_path.name}")
        print("\n💡 使用提示:")
        print("  • PDF开头是图片索引目录（A3纵向，支持多页）")
        print(f"  • 有效图片: 连续编号 图1~图{active_count}")
        if orphan_count > 0:
            print("  • 孤儿图片: PDF末尾 (标注[未使用])")
        return True
    except Exception as e:
        print(f"❌ 处理失败: {e}")
        traceback.print_exc()
        return False


# --- 模式编排与内置功能集成 ---

def is_aio_processing_needed(args):
    """是否启用AIO核心处理模式（Excel对象处理 + 文档优化）"""
    return any([
        getattr(args, 'word_table', False),
        getattr(args, 'extract_excel', False),
        getattr(args, 'image', False),
        getattr(args, 'remove_watermark', False),
        getattr(args, 'a3', False),
    ])


def build_aio_output_path(input_path, args):
    """根据参数生成AIO核心输出路径"""
    base, ext = os.path.splitext(input_path)
    suffix_parts = ["-AIO"]
    if getattr(args, 'keep_attachment', False):
        suffix_parts.append("WithAttachments")
    if getattr(args, 'remove_watermark', False):
        suffix_parts.append("NoWM")
    if getattr(args, 'a3', False):
        suffix_parts.append("A3")

    suffix = "-" + "-".join(suffix_parts[1:]) if len(suffix_parts) > 1 else suffix_parts[0]
    return f"{base}{suffix}{ext}"


def _check_extension_deps(need_reportlab=True, need_pil=False):
    """检查扩展功能所需的可选依赖"""
    missing = []
    if need_reportlab and not _HAS_REPORTLAB:
        missing.append("reportlab")
    if need_pil and not _HAS_PIL:
        missing.append("Pillow")
    if missing:
        deps = " ".join(missing)
        print(f"❌ 缺少依赖: {deps}")
        print(f"   请安装: pip install {deps}")
        raise RuntimeError(f"缺少依赖: {deps}")


def run_table_extract_extension(target_docx_path, args):
    """执行内置表格提取流程"""
    _check_extension_deps(need_reportlab=True, need_pil=False)
    result = process_docx(target_docx_path)
    if result is False:
        raise RuntimeError("表格提取失败")
    if result is None:
        print("    ⏭️  表格提取跳过: 文档中没有可提取表格")


def run_image_split_extension(target_docx_path, args):
    """执行内置图片分离流程"""
    _check_extension_deps(need_reportlab=True, need_pil=True)
    ok = process_docx_file(
        target_docx_path,
        remove_images=getattr(args, 'split_remove_images', False),
        output_dir=getattr(args, 'split_output_dir', None),
        optimize_images=not getattr(args, 'split_no_optimize', False),
        jpeg_quality=getattr(args, 'split_jpeg_quality', 85),
    )
    if not ok:
        raise RuntimeError("图片分离失败")


def process_document_with_extensions(input_path, args):
    """统一编排：先执行AIO核心功能，再按需执行内置附加功能"""
    target_docx_path = input_path

    if is_aio_processing_needed(args):
        process_document(input_path, args)
        target_docx_path = build_aio_output_path(input_path, args)
    else:
        print("ℹ️  未启用AIO核心模式，直接执行扩展功能。")

    if getattr(args, 'table_extract', False):
        print(f"\n🧩 执行附加功能: 表格提取 -> {target_docx_path}")
        run_table_extract_extension(target_docx_path, args)

    if getattr(args, 'split_images', False):
        print(f"\n🧩 执行附加功能: 图片分离 -> {target_docx_path}")
        run_image_split_extension(target_docx_path, args)


# --- 增强版主处理函数 ---

def process_document(input_path, args):
    """智能文档处理主函数

    功能说明:
        两阶段处理DOCX文档中的嵌入式Excel对象，支持多种转换模式和优化选项。

    处理流程:
        【准备阶段】
        1. 根据参数生成输出文件名（-AIO、-WithAttachments、-NoWM、-A3后缀）
        2. 使用EnhancedExcelLocator分析文档结构
        3. 提取所有嵌入的Excel数据（从word/embeddings目录）
        4. 定位Excel对象和说明文字段落

        【第一阶段 - 内容插入】
        5. 应用页面设置（A3横向）和水印清理（如启用）
        6. 从后往前遍历Excel对象（避免索引错位）
        7. 根据启用的模式插入内容:
           - extract-excel: 提取.xlsx文件 + 添加"表 X"标题
           - image:         渲染matplotlib图片 + 插入到文档
           - word-table:    转换为Word原生表格 + 保留完整样式
        8. 保存到临时文件

        【第二阶段 - 智能清理】
        9. 重新加载临时文件
        10. 扫描并标记需要清理的元素:
            - Excel对象段落（如未启用--keep-attachment）
            - 说明文字段落（占位符文本）
        11. 执行删除操作（从DOM树移除）
        12. 保存最终文档并清理临时文件

        【验证和报告】
        13. 统计处理结果（表格数、文件数、大小变化）
        14. 输出详细的处理报告

    支持的模式组合:
        - word-table + extract-excel + image: 全功能模式
        - word-table + keep-attachment:       保留附件可点击
        - image + a3:                         大表格图片优化
        - 任意组合都是允许的

    Args:
        input_path (str): 输入DOCX文件的完整路径
        args (Namespace): argparse解析的命令行参数对象，包含:
            - word_table (bool):      是否转换为Word表格
            - extract_excel (bool):   是否提取Excel文件
            - image (bool):           是否渲染为图片
            - keep_attachment (bool): 是否保留Excel附件
            - remove_watermark (bool): 是否移除水印
            - a3 (bool):              是否使用A3横向布局

    输出文件命名:
        基础: [原文件名]-AIO.docx
        后缀:
            - WithAttachments: 保留了Excel附件
            - NoWM:           移除了水印
            - A3:             A3横向布局
        示例: report-AIO-A3-NoWM.docx

    关键技术:
        - 逆序处理: 从后往前遍历，避免删除时索引变化
        - 段落保护: 检测段落父节点，不处理表格内的对象
        - 两阶段操作: 先插入后删除，确保文档结构完整性
        - 样式保留: 完整保留Excel的字体、颜色、对齐、合并等样式

    异常处理:
        - 如果没有找到Excel数据，打印错误并退出
        - 如果没有找到Excel对象位置，打印错误并退出
        - 单个处理步骤失败时打印警告，继续处理其他对象
    """

    base, ext = os.path.splitext(input_path)
    output_path = build_aio_output_path(input_path, args)
    
    print(f"🚀 增强版处理: {input_path}")
    print(f"    📁 输出: {output_path}")
    print(f"    🔗 保留附件: {'是' if args.keep_attachment else '否'}")
    print(f"    🧹 移除水印: {'是' if getattr(args, 'remove_watermark', False) else '否'}")
    print(f"    📄 A3横向: {'是' if getattr(args, 'a3', False) else '否'}")
    
    # 分析文档
    locator = EnhancedExcelLocator(input_path)
    doc_info = locator.analyze_document_structure()
    
    if doc_info['original_tables'] > 0:
        print(f"    🛡️  检测到 {doc_info['original_tables']} 个原始表格，将完全保护")
    
    # 提取Excel数据
    all_excels = extract_embedded_excel_enhanced(input_path)

    # 定位Excel对象
    excel_objects, caption_paragraphs = [], []
    if all_excels:
        excel_objects, caption_paragraphs = locator.find_excel_objects_enhanced()

    # 检查是否有Excel相关的处理需求
    has_excel_mode = args.word_table or args.extract_excel or args.image
    has_other_mode = (
        getattr(args, 'remove_watermark', False) or
        getattr(args, 'a3', False) or
        getattr(args, 'table_extract', False) or
        getattr(args, 'split_images', False)
    )

    if not all_excels or not excel_objects:
        if has_excel_mode and not has_other_mode:
            # 只有Excel模式但没有Excel数据，跳过
            print("❌ 未找到Excel数据，且未启用其他处理模式")
            raise DocumentSkipError("未找到Excel数据")
        elif has_excel_mode:
            # 有Excel模式但也有其他模式，警告但继续
            print("⚠️  未找到Excel数据，将跳过Excel相关处理，继续执行其他功能")
        # 如果只有其他模式（remove_watermark, a3），则继续处理

    print(f"\n📋 处理计划:")
    print(f"    Excel对象: {len(excel_objects)} 个")
    print(f"    说明文字: {len(caption_paragraphs)} 个")
    if has_excel_mode and excel_objects:
        print(f"    处理模式: {'转换表格+保留附件' if args.keep_attachment else '转换表格+完全清理'}")
    
    # 第一阶段：处理所有模式
    doc = Document(input_path)
    processed_count = 0
    
    # 应用页面设置和清理功能
    if getattr(args, 'a3', False):
        setup_a3_landscape_page(doc)
    
    if getattr(args, 'remove_watermark', False):
        cleaner = DocumentCleaner(doc)
        cleaner.remove_watermarks()
    
    # ============================================================
    # 第一阶段处理逻辑：从后往前遍历Excel对象
    # ============================================================
    #
    # 逆序遍历的原因：
    # 1. 避免索引错位: 当删除或修改前面的段落时，后面段落的索引会变化
    # 2. 保持文档顺序: 从后往前处理，插入的内容在文档中仍然保持原有顺序
    # 3. 安全性: 即使某个对象处理失败，也不会影响其他对象的索引定位
    #
    # 示例：文档中有3个Excel对象在段落 [5, 10, 15]
    # - 正序处理: 处理段落5后，原段落10变成9，原段落15变成14（索引错位）
    # - 逆序处理: 先处理15，再处理10，最后处理5（索引始终准确）
    for i in range(min(len(excel_objects), len(all_excels)) - 1, -1, -1):
        excel_obj = excel_objects[i]
        excel_data, excel_name = all_excels[i]
        excel_idx = excel_obj['index']

        print(f"\n    🎯 处理Excel对象 {i+1}: 段落 {excel_idx+1}")

        placeholder = doc.paragraphs[excel_idx] if excel_idx < len(doc.paragraphs) else None

        # ============================================================
        # 内部辅助函数：智能元素插入
        # ============================================================
        #
        # 功能：根据占位符是否存在，选择不同的插入策略
        #
        # 策略1：占位符存在（正常情况）
        #   - 使用addprevious()在占位符段落之前插入新元素
        #   - 保持文档结构清晰，新内容紧挨着原Excel对象
        #
        # 策略2：占位符不存在（异常情况）
        #   - 直接append到document.body末尾
        #   - 确保内容不会丢失，即使定位失败
        #
        # 参数说明：
        #   element: 要插入的XML元素（表格._tbl 或 段落._p）
        #   is_table: 是否为表格元素（影响XML结构）
        def insert_element(element, is_table=False):
            if placeholder:
                # 表格使用_tbl，段落等使用_p
                target_p = placeholder._p
                target_p.addprevious(element)
            else:
                # 如果是表格，需要添加到body
                if is_table:
                    doc.element.body.append(element)

        # ============================================================
        # 模式处理顺序说明
        # ============================================================
        #
        # 处理顺序：Word表格 -> 图片 -> Excel提取
        #
        # 为什么是这个顺序？
        # 1. 使用addprevious()插入时，后插入的元素会出现在前面
        # 2. 期望的文档顺序: [提取标题] -> [图片] -> [Word表格] -> [原Excel对象]
        # 3. 因此代码执行顺序要反过来: Word表格 -> 图片 -> 提取
        #
        # 举例：
        #   执行: create_table() -> add_picture() -> add_paragraph("表 1")
        #   结果: [表 1] -> [图片] -> [Word表格] -> [原Excel对象]
        #   这样阅读顺序是正确的！
        
        # 模式3：插入为Word表格
        if args.word_table:
            excel_data.seek(0)
            data, merges, rows, cols = get_excel_data_enhanced(excel_data, excel_name)
            if data and rows > 0 and cols > 0:
                new_table = create_table_enhanced(
                    doc, data, merges, rows, cols, 
                    a3_landscape=getattr(args, 'a3', False)
                )
                insert_element(new_table._tbl, is_table=True)
                processed_count += 1
                print(f"        ✅ Word表格已插入")
            else:
                print(f"        ❌ Word表格转换失败")
        
        # 模式2：插入为图片
        if args.image:
            temp_img = tempfile.NamedTemporaryFile(
                prefix=f"docx_aio_{os.getpid()}_{i}_",
                suffix=".png",
                delete=False
            )
            img_path = temp_img.name
            temp_img.close()

            excel_data.seek(0)
            try:
                if excel_to_image(excel_data, img_path, a3_landscape=getattr(args, 'a3', False)):
                    print(f"        ✅ 已使用 matplotlib 将表格渲染为图片")
                    if placeholder:
                        p = placeholder.insert_paragraph_before()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        # A3横向时使用更大的图片宽度
                        img_width = Inches(10.0) if getattr(args, 'a3', False) else Inches(6.5)
                        run.add_picture(img_path, width=img_width)
                    else:
                        img_width = Inches(10.0) if getattr(args, 'a3', False) else Inches(6.5)
                        doc.add_picture(img_path, width=img_width)
                else:
                    print(f"        ❌ 图片渲染失败")
            finally:
                try:
                    if os.path.exists(img_path):
                        os.remove(img_path)
                except Exception as e:
                    print(f"        ⚠️  临时图片清理失败: {e}")
        
        # 模式1：提取Excel文件（并添加标题）
        if args.extract_excel:
            base_name, _ = os.path.splitext(input_path)
            excel_path = f"{base_name}_table_{i+1}.xlsx"
            try:
                with open(excel_path, 'wb') as f:
                    excel_data.seek(0)
                    f.write(excel_data.read())
                print(f"        ✅ 已提取Excel到: {excel_path}")
                
                caption = f"表 {i+1}"
                if placeholder:
                    p = placeholder.insert_paragraph_before(caption)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # 如果没有占位符，在追加内容前添加标题
                    doc.add_paragraph(caption, style='Body Text').alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"        ❌ 提取Excel失败: {e}")
    
    # 保存临时文件
    temp_path = f"{base}-temp{ext}"
    doc.save(temp_path)
    print(f"    💾 临时保存完成")
    
    # ============================================================
    # 第二阶段：智能清理（从临时文件重新加载）
    # ============================================================
    #
    # 为什么需要第二阶段？
    # 1. 分离关注点: 第一阶段专注插入新内容，第二阶段专注清理旧内容
    # 2. 数据完整性: 先保存临时文件，确保新内容已经写入
    # 3. 安全操作: 重新加载后再删除，避免XML结构混乱
    # 4. 灵活控制: 可以根据--keep-attachment参数决定是否删除原Excel对象
    #
    # 清理目标：
    # - Excel嵌入对象段落（如果未启用--keep-attachment）
    # - 说明文字段落（"点击图片可查看完整电子表格"等占位符）
    doc = Document(temp_path)

    excel_objects_to_remove = []
    captions_to_remove = []

    # 扫描文档，标记需要清理的段落
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_xml = ET.tostring(paragraph._p, encoding='unicode')
        para_text = paragraph.text.strip()

        # ============================================================
        # Excel对象识别（简化版）
        # ============================================================
        #
        # 注意：这里的识别逻辑比第一阶段简单，因为：
        # 1. 不需要匹配关系ID（已经提取过数据）
        # 2. 只需要识别基本的VML/OLE结构特征
        # 3. 重点是准确定位，避免误删
        has_vml_shape = ('<v:shape' in para_xml or '<ns2:shape' in para_xml) and ('ole="t"' in para_xml)
        has_ole_object = '<o:OLEObject' in para_xml or '<ns3:OLEObject' in para_xml
        has_excel_progid = 'ProgID="Excel.Sheet.12"' in para_xml
        is_in_table = locator._is_paragraph_in_table(paragraph, doc)

        # ============================================================
        # Excel对象清理决策
        # ============================================================
        #
        # 清理条件：
        # 1. 包含VML shape或OLE对象
        # 2. 包含Excel ProgID
        # 3. 不在现有表格中
        #
        # 清理模式：
        # - 默认模式（--keep-attachment未启用）: 完全删除Excel对象
        # - 保留模式（--keep-attachment启用）: 保留Excel对象，用户仍可双击打开
        if ((has_vml_shape or has_ole_object) and has_excel_progid and not is_in_table):
            if not args.keep_attachment:
                # 默认模式：完全清理Excel对象
                excel_objects_to_remove.append(paragraph)
                print(f"    🗑️  标记删除Excel对象: 段落 {para_idx}")
            else:
                # 保留附件模式：Excel对象保留
                print(f"    🔗 保留Excel附件: 段落 {para_idx}")

        # ============================================================
        # 说明文字清理
        # ============================================================
        #
        # 说明文字：嵌入Excel对象时Word自动添加的占位符文本
        # 常见内容："点击图片可查看完整电子表格"、"Click to view complete spreadsheet"
        #
        # 清理策略：
        # 1. 精确匹配CAPTION_KEYWORDS中的关键词
        # 2. 确保不在现有表格中（避免误删用户内容）
        # 3. 无论是否保留附件，说明文字都应该清理
        if para_text in CAPTION_KEYWORDS and not is_in_table:
            # 清理说明文字
            captions_to_remove.append(paragraph)
            print(f"    🗑️  标记删除说明文字: 段落 {para_idx}")
    
    # 执行删除操作
    removed_count = 0
    
    for para in excel_objects_to_remove:
        try:
            p_element = para._p
            if p_element.getparent() is not None:
                p_element.getparent().remove(p_element)
                removed_count += 1
        except Exception as e:
            print(f"    ⚠️  删除Excel对象失败: {e}")
    
    for para in captions_to_remove:
        try:
            p_element = para._p
            if p_element.getparent() is not None:
                p_element.getparent().remove(p_element)
                removed_count += 1
        except Exception as e:
            print(f"    ⚠️  删除说明文字失败: {e}")
    
    # 保存最终文档
    doc.save(output_path)
    
    # 清理临时文件
    try:
        if os.path.exists(temp_path):
            os.remove(temp_path)
    except Exception as e:
        print(f"    ⚠️  临时文档清理失败: {e}")
    
    # 验证结果
    final_doc = Document(output_path)
    final_tables = len(final_doc.tables)
    
    print(f"\n🎉 全功能处理完成!")
    print(f"✅ 文件已保存: {output_path}")
    print(f"📊 处理结果:")
    # 计算处理模式
    modes = []
    if args.word_table: modes.append("Word表格")
    if args.extract_excel: modes.append("提取Excel")
    if args.image: modes.append("图片渲染")
    if args.keep_attachment: modes.append("保留附件")
    
    print(f"    处理模式: {', '.join(modes)}")
    print(f"    处理Excel对象: {processed_count} 个")
    print(f"    删除段落: {removed_count} 个")
    print(f"    原始表格: {doc_info['original_tables']} 个 (完全保护)")
    print(f"    最终表格数: {final_tables} 个")
    if args.extract_excel:
        print(f"    提取文件: {processed_count} 个 .xlsx")
    if args.image:
        print(f"    渲染图片: {processed_count} 个 .png")
    print(f"    附件保留: {'是' if args.keep_attachment else '否'}")
    
    # 文件大小
    original_size = os.path.getsize(input_path)
    output_size = os.path.getsize(output_path)
    print(f"📏 文件大小: {original_size:,} → {output_size:,} bytes ({((output_size-original_size)/original_size*100):+.1f}%)")


def get_docx_files_from_folder(folder_path):
    """获取文件夹中需要处理的docx文件列表

    过滤规则：
    1. 只处理 .docx 文件
    2. 跳过文件名包含输出标签的文件（-WithAttachments, -NoWM, -A3, -AIO）
    3. 不处理子文件夹中的文件
    4. 跳过临时文件（以 ~$ 开头）

    Args:
        folder_path: 文件夹路径

    Returns:
        list: 需要处理的docx文件完整路径列表
    """
    docx_files = []

    # 获取文件夹中的所有文件（不包含子文件夹）
    try:
        for item in os.listdir(folder_path):
            item_path = os.path.join(folder_path, item)

            # 跳过子文件夹
            if os.path.isdir(item_path):
                continue

            # 只处理 .docx 文件
            if not item.lower().endswith('.docx'):
                continue

            # 跳过临时文件
            if item.startswith('~$'):
                continue

            # 检查是否包含输出标签
            file_base = os.path.splitext(item)[0]
            has_output_tag = any(tag in file_base for tag in OUTPUT_FILE_TAGS)

            if has_output_tag:
                print(f"    ⏭️  跳过已处理文件: {item}")
                continue

            docx_files.append(item_path)

    except Exception as e:
        print(f"❌ 读取文件夹失败: {e}")
        return []

    return sorted(docx_files)  # 按文件名排序


def process_batch(folder_path, args):
    """批量处理文件夹中的docx文件

    Args:
        folder_path: 文件夹路径
        args: 命令行参数
    """
    print(f"📂 批量处理模式: {folder_path}")
    print(f"    🔍 扫描文件夹...")

    docx_files = get_docx_files_from_folder(folder_path)

    if not docx_files:
        print("❌ 未找到需要处理的docx文件")
        return

    print(f"    📄 找到 {len(docx_files)} 个待处理文件\n")

    workers = min(max(1, getattr(args, 'workers', 1)), len(docx_files))

    # workers=1：保持原始日志行为
    if workers == 1:
        success_count, fail_count, skip_count = process_batch_serial(docx_files, args)
    else:
        success_count, fail_count, skip_count = process_batch_parallel(docx_files, args, workers)

    # 批量处理总结
    print(f"\n{'='*60}")
    print(f"📊 批量处理完成!")
    print(f"{'='*60}")
    print(f"    ✅ 成功: {success_count} 个")
    print(f"    ⏭️  跳过: {skip_count} 个")
    print(f"    ❌ 失败: {fail_count} 个")
    print(f"    📁 总计: {len(docx_files)} 个文件")


def process_batch_serial(docx_files, args):
    """串行批处理（保留当前输出风格）"""
    success_count = 0
    fail_count = 0
    skip_count = 0

    for idx, file_path in enumerate(docx_files, 1):
        file_name = os.path.basename(file_path)
        print(f"\n{'='*60}")
        print(f"📄 [{idx}/{len(docx_files)}] 处理文件: {file_name}")
        print(f"{'='*60}")

        try:
            process_document_with_extensions(file_path, args)
            success_count += 1
        except DocumentSkipError as e:
            print(f"⚠️  文件处理跳过: {file_name} ({e})")
            skip_count += 1
        except Exception as e:
            print(f"❌ 文件处理失败: {file_name}")
            print(f"    错误信息: {e}")
            fail_count += 1

    return success_count, fail_count, skip_count


def extract_core_log_lines(log_output):
    """提取多worker场景下需要展示的核心日志"""
    lines = [line.rstrip() for line in log_output.splitlines() if line.strip()]
    core_lines = []

    # 保留输出路径
    output_line = next((line.strip() for line in lines if "✅ 文件已保存:" in line), None)
    if output_line:
        core_lines.append(output_line)

    # 保留最终统计核心信息
    in_result_block = False
    stat_keys = ("处理模式:", "处理Excel对象:", "删除段落:", "提取文件:", "渲染图片:", "附件保留:")
    for line in lines:
        stripped = line.strip()
        if stripped == "📊 处理结果:":
            in_result_block = True
            continue
        if in_result_block:
            if line.startswith("    "):
                if any(key in stripped for key in stat_keys):
                    core_lines.append(stripped)
                continue
            in_result_block = False

    # 额外保留关键告警/错误（最多3条）
    warn_or_error = [line.strip() for line in lines if ("❌" in line or "⚠️" in line)]
    for item in warn_or_error[:3]:
        if item not in core_lines:
            core_lines.append(item)

    if not core_lines:
        core_lines = [line.strip() for line in lines[-3:]]

    return core_lines[:10]


def process_document_worker(file_path, args_dict):
    """多进程worker：捕获完整日志并回传核心摘要"""
    args = argparse.Namespace(**args_dict)
    output_buffer = io.StringIO()

    status = "success"
    error_message = ""

    try:
        with contextlib.redirect_stdout(output_buffer), contextlib.redirect_stderr(output_buffer):
            process_document_with_extensions(file_path, args)
    except DocumentSkipError as e:
        status = "skip"
        error_message = str(e)
    except Exception as e:
        status = "fail"
        error_message = f"{e}\n{traceback.format_exc(limit=6)}"

    full_log = output_buffer.getvalue()

    return {
        "status": status,
        "pid": os.getpid(),
        "file_path": file_path,
        "core_logs": extract_core_log_lines(full_log),
        "error": error_message,
    }


def process_batch_parallel(docx_files, args, workers):
    """并行批处理（输出核心日志，避免交错）"""
    total = len(docx_files)
    success_count = 0
    fail_count = 0
    skip_count = 0

    print(f"    🧵 并行worker: {workers}")
    print(f"    🏷️  日志模式: 核心摘要 + 标记输出")

    args_dict = vars(args).copy()
    future_map = {}

    with concurrent.futures.ProcessPoolExecutor(max_workers=workers) as executor:
        for idx, file_path in enumerate(docx_files, 1):
            file_name = os.path.basename(file_path)
            print(f"[QUEUE][{idx}/{total}] {file_name}")
            future = executor.submit(process_document_worker, file_path, args_dict)
            future_map[future] = (idx, file_name)

        completed = 0
        for future in concurrent.futures.as_completed(future_map):
            completed += 1
            idx, file_name = future_map[future]

            try:
                result = future.result()
            except Exception as e:
                fail_count += 1
                print(f"[FAIL][{completed}/{total}][#{idx}] {file_name}")
                print(f"    [ERR] worker异常: {e}")
                continue

            status = result.get("status", "fail")
            pid = result.get("pid", "-")
            core_logs = result.get("core_logs", [])
            error_message = result.get("error", "")

            if status == "success":
                success_count += 1
                status_tag = "OK"
            elif status == "skip":
                skip_count += 1
                status_tag = "SKIP"
            else:
                fail_count += 1
                status_tag = "FAIL"

            print(f"[{status_tag}][{completed}/{total}][#{idx}][PID:{pid}] {file_name}")
            for line in core_logs:
                print(f"    [CORE] {line}")

            if error_message and status != "success":
                print(f"    [ERR] {error_message}")

    return success_count, fail_count, skip_count


def main():
    parser = argparse.ArgumentParser(
        description="一个多功能CLI工具，用于处理DOCX中的嵌入式Excel表格。",
        formatter_class=argparse.RawTextHelpFormatter,
        epilog="""🚀 全功能特性:
• 支持多模式组合使用
• 支持单文件或文件夹批量处理
• 自适应表格优化：列宽和行高自动调整(1.5-6.0cm)
• 智能图片渲染：中文字体支持和高质量输出(300DPI)
• 完全保护原始Word表格
• 精准Excel对象识别和处理
• 智能水印移除：支持文本、图片、背景等多种水印格式
• A3横向布局：专为大表格优化的页面设置
• 单文件内置：表格提取与图片分离能力已内联

使用模式:
# 默认模式 (转为Word表格)
python %(prog)s my_document.docx

# 将表格转为图片
python %(prog)s my_document.docx --image

# 组合使用：提取Excel，并转为图片插入
python %(prog)s my_document.docx --extract-excel --image

# 三种模式全开启（转Word表格+提取Excel+生成图片）
python %(prog)s my_document.docx --word-table --extract-excel --image

# 保留附件模式组合
python %(prog)s my_document.docx --word-table --keep-attachment

# 移除水印并设置A3横向
python %(prog)s my_document.docx --remove-watermark --a3

# A3横向布局 + 图片渲染
python %(prog)s my_document.docx --image --a3

# 全功能组合：Word表格 + 图片 + A3横向 + 无水印
python %(prog)s my_document.docx --word-table --image --a3 --remove-watermark

# 只做表格提取（输出TXT/XLSX/PDF+标记文档）
python %(prog)s my_document.docx --table-extract

# 只做图片分离（输出附图PDF+图片标记文档）
python %(prog)s my_document.docx --split-images

# 先AIO转换，再做表格提取+图片分离
python %(prog)s my_document.docx --word-table --remove-watermark --table-extract --split-images

# 📂 批量处理文件夹（自动跳过已处理文件）
python %(prog)s /path/to/folder --remove-watermark --a3

# 批量处理：转换所有docx为Word表格
python %(prog)s /path/to/folder --word-table
"""
    )
    parser.add_argument("input_path", help="输入DOCX文件路径或包含DOCX文件的文件夹路径。")
    
    # 模式参数
    parser.add_argument(
        "--word-table",
        action="store_true",
        help="将Excel表格转换为可编辑的Word原生表格。"
    )
    parser.add_argument(
        "--extract-excel",
        action="store_true",
        help="提取嵌入的Excel文件为独立.xlsx文件，并在原位置标注'表 X'。"
    )
    parser.add_argument(
        "--image",
        action="store_true",
        help="将Excel表格渲染成图片（使用matplotlib引擎）。"
    )
    parser.add_argument(
        "--keep-attachment", 
        action="store_true", 
        help="保留Excel附件入口，用户仍可点击查看完整表格。"
    )
    
    # 新增功能参数
    parser.add_argument(
        "--remove-watermark",
        action="store_true",
        help="移除文档水印，包括页眉页脚和正文中的水印元素。"
    )
    parser.add_argument(
        "--a3",
        action="store_true", 
        help="设置文档为A3横向页面 (420mm×297mm)，优化表格和图片布局。"
    )

    # 内置附加模式（单文件实现）
    parser.add_argument(
        "--table-extract",
        action="store_true",
        help="内置表格提取：导出表格到 TXT/XLSX/PDF，并生成已标记表格文档。"
    )
    parser.add_argument(
        "--split-images",
        action="store_true",
        help="内置图片分离：导出附图PDF，并生成已标记图片文档。"
    )
    parser.add_argument(
        "--split-remove-images",
        action="store_true",
        help="配合 --split-images 使用：在标记文档中删除原图片，仅保留图号标记。"
    )
    parser.add_argument(
        "--split-output-dir",
        help="配合 --split-images 使用：指定图片分离输出目录（默认与目标DOCX同目录）。"
    )
    parser.add_argument(
        "--split-no-optimize",
        action="store_true",
        help="配合 --split-images 使用：关闭图片优化（输出体积更大）。"
    )
    parser.add_argument(
        "--split-jpeg-quality",
        type=int,
        default=85,
        metavar="Q",
        help="配合 --split-images 使用：JPEG质量(1-100)，默认85。"
    )
    parser.add_argument(
        "--workers",
        type=int,
        default=1,
        help="批量处理worker数量。1=串行并保持原打印；>1=并行并输出核心标记日志。"
    )
    
    if len(sys.argv) == 1:
        parser.print_help(sys.stderr)
        sys.exit(1)
        
    args = parser.parse_args()

    if args.workers < 1:
        print("❌ 错误: --workers 必须大于等于 1。")
        sys.exit(1)

    if not 1 <= args.split_jpeg_quality <= 100:
        print("❌ 错误: --split-jpeg-quality 必须在 1~100 之间。")
        sys.exit(1)

    if not args.split_images and (
        args.split_remove_images or args.split_output_dir or args.split_no_optimize or args.split_jpeg_quality != 85
    ):
        print("❌ 错误: --split-remove-images/--split-output-dir/--split-no-optimize/--split-jpeg-quality 需要和 --split-images 一起使用。")
        sys.exit(1)

    # 检查是否有任何处理模式
    has_excel_mode = any([args.word_table, args.extract_excel, args.image])
    has_other_mode = getattr(args, 'remove_watermark', False) or getattr(args, 'a3', False)
    has_extension_mode = args.table_extract or args.split_images

    # 如果没有指定任何模式，则默认使用 --word-table
    if not has_excel_mode and not has_other_mode and not has_extension_mode:
        args.word_table = True
        print("[i] 未指定任何操作模式，将默认执行 --word-table。")

    # 检查是否至少选择了一种操作模式
    has_any_mode = (
        args.word_table or args.extract_excel or args.image or args.remove_watermark or args.a3 or
        args.table_extract or args.split_images
    )
    if not has_any_mode:
        print("❌ 错误: 必须至少选择一种操作模式。")
        print("    Excel模式: --word-table, --extract-excel, --image")
        print("    文档优化: --remove-watermark, --a3")
        print("    附加模式: --table-extract, --split-images")
        sys.exit(1)

    if not os.path.exists(args.input_path):
        print(f"❌ 错误: 输入路径不存在 -> {args.input_path}")
        sys.exit(1)

    # 判断输入是文件还是文件夹
    if os.path.isdir(args.input_path):
        # 文件夹批量处理模式
        process_batch(args.input_path, args)
    elif os.path.isfile(args.input_path):
        # 单文件处理模式
        if not args.input_path.lower().endswith('.docx'):
            print(f"❌ 错误: 输入文件必须是 .docx 格式 -> {args.input_path}")
            sys.exit(1)
        process_document_with_extensions(args.input_path, args)
    else:
        print(f"❌ 错误: 无法识别的输入类型 -> {args.input_path}")
        sys.exit(1)


if __name__ == "__main__":
    main()
